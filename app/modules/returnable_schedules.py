"""
returnable_schedules.py

Returnable-schedule filler: takes the client's own response forms (DOCX
schedules built as tables/label-value forms, XLSX schedules) and fills them
from what this project actually knows -- company/contact details, the
resourcing plan's personnel, reference projects, fee build-up lines --
INSIDE the client's original file, preserving their formatting. That last
part is the whole point: councils and principals reject re-typed schedules;
the filled document must be their document.

The no-invention rule is absolute here, same as everywhere else in this
app: a field the project genuinely knows gets its real value; everything
else gets the standard clearly-marked placeholder
("[TO BE COMPLETED: <what>]") -- never a guess. Insurance details, ABN/ACN,
and anything commercial the user hasn't entered are deliberately in the
"never known" bucket.

Filling is deterministic (synonym/keyword matching), not AI-driven, on
purpose: an AI mapper could hallucinate a mapping and put a *wrong real
value* in a legal form, which is worse than a placeholder. A label that
can't be confidently matched gets a placeholder naming the label, so the
user can't miss it.

Entry points:
  looks_like_response_form(filename, file_bytes) -> bool
  build_fill_data(state)                         -> dict (see FILL_FIELDS)
  fill_schedule(filename, file_bytes, fill_data) -> FillResult
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field

PLACEHOLDER_PREFIX = "[TO BE COMPLETED"


def make_placeholder(label: str) -> str:
    label = re.sub(r"\s+", " ", (label or "").strip().strip(":*")).strip()
    return f"{PLACEHOLDER_PREFIX}: {label}]" if label else f"{PLACEHOLDER_PREFIX}]"


# ---------------------------------------------------------------------------
# What the project can know, and the label synonyms that map to each field.
# Order matters where synonyms overlap -- more specific entries first.
# ---------------------------------------------------------------------------

# Firm-level labels a returnable schedule asks about. These used to be in
# the NEVER_KNOWN bucket below, unconditionally -- correct while the app had
# nowhere to store an ABN or an insurance policy, and wrong the moment the
# firm profile existed. They are matched BEFORE the never-known patterns and
# only answered when the profile actually holds a value; with an empty
# profile the never-known behaviour is unchanged, placeholder and all.
#
# "signature" is deliberately absent and always will be: nobody should be
# able to sign a legal form out of a saved setting.
FIRM_FIELD_SYNONYMS: list[tuple[str, list[str]]] = [
    ("professional_indemnity_insurer", ["professional indemnity insurer", "pi insurer"]),
    ("professional_indemnity_policy", ["professional indemnity policy", "pi policy"]),
    ("professional_indemnity_cover", ["professional indemnity cover", "professional indemnity limit",
                                       "professional indemnity amount", "pi cover"]),
    ("professional_indemnity_expiry", ["professional indemnity expiry", "pi expiry"]),
    ("public_liability_insurer", ["public liability insurer"]),
    ("public_liability_policy", ["public liability policy"]),
    ("public_liability_cover", ["public liability cover", "public liability limit",
                                 "public liability amount"]),
    ("public_liability_expiry", ["public liability expiry"]),
    ("workers_compensation_insurer", ["workers compensation insurer", "workers comp insurer"]),
    ("workers_compensation_policy", ["workers compensation policy", "workers comp policy"]),
    ("workers_compensation_cover", ["workers compensation cover", "workers comp cover"]),
    ("workers_compensation_expiry", ["workers compensation expiry", "workers comp expiry"]),
    ("company_abn", ["abn", "australian business number"]),
    ("company_acn", ["acn", "australian company number"]),
    ("certifications", ["certification", "accreditation", "quality certification",
                        "iso certification", "quality accreditation"]),
]

FIELD_SYNONYMS: list[tuple[str, list[str]]] = [
    ("contact_email",   ["email address", "e-mail", "email"]),
    ("contact_phone",   ["phone number", "telephone", "phone", "mobile", "contact number"]),
    ("contact_name",    ["contact person", "contact name", "authorised representative",
                         "authorised person", "nominated contact", "contact officer",
                         "name of authorised", "signatory", "signed by", "name of person"]),
    ("contact_title",   ["position of contact", "position/title", "position", "title of signatory"]),
    ("contact_address", ["registered office", "business address", "postal address",
                         "street address", "address"]),
    ("company_name",    ["name of tenderer", "tenderer's name", "name of respondent",
                         "legal entity name", "entity name", "trading name", "company name",
                         "business name", "name of supplier", "name of contractor",
                         "name of consultant", "organisation name", "organization name",
                         "tenderer", "respondent", "supplier", "consultant name"]),
    ("tender_number",   ["tender number", "tender no", "rft number", "rft no", "rfq number",
                         "rfq no", "contract number", "contract no", "reference number",
                         "tender reference", "rft ref"]),
    ("project_name",    ["project title", "contract title", "project name", "name of project",
                         "title of contract", "contract name", "project description"]),
    ("client_name",     ["name of principal", "principal", "client name", "client"]),
    ("submission_date", ["closing date", "date of submission", "submission date", "date of tender", "date"]),
]

# Fields the app never knows unless the user has typed them somewhere it
# doesn't currently capture -- ALWAYS placeholdered, with the label kept so
# the user knows exactly what to complete. Matching any of these takes
# priority over the generic label fallback (same outcome, clearer intent).
NEVER_KNOWN_PATTERNS = [
    r"\babn\b", r"australian business number", r"\bacn\b", r"australian company number",
    r"insur", r"indemnity", r"public liability", r"workers'? comp", r"policy",
    r"premium", r"\bgst\b", r"bank", r"\bbsb\b", r"account", r"licen[cs]e",
    r"\bqbcc\b", r"registration number", r"signature", r"\bsign\b",
]

# Roster-table column headers -> personnel attributes (see _fill_roster_table).
ROSTER_COLUMNS = {
    "name": ["name", "personnel", "person", "team member", "employee"],
    "role": ["role", "position", "title", "discipline", "function", "responsibility"],
    "quals": ["qualification", "quals", "accreditation", "registration", "rpeq", "certification"],
    "years": ["years", "experience"],
}

REFERENCE_COLUMNS = {
    "title": ["project", "contract", "title", "description of project", "name of project"],
    "client": ["client", "principal", "for whom"],
    "description": ["description", "scope", "details", "relevance"],
}


@dataclass
class FillResult:
    filename: str
    file_bytes: bytes | None          # the filled file; None when nothing could be done
    kind: str                          # "docx" | "xlsx" | "unsupported"
    filled: list[dict] = field(default_factory=list)        # {"where", "label", "value"}
    placeholdered: list[dict] = field(default_factory=list) # {"where", "label"}
    error: str | None = None


# ---------------------------------------------------------------------------
# Fill data -- everything the project knows, in one flat dict + lists
# ---------------------------------------------------------------------------

def build_fill_data(state, firm_data: dict | None = None) -> dict:
    """`state` is anything with .get() -- in practice st.session_state.
    Only REAL user-entered values end up here; blank stays absent, so the
    filler placeholders it.

    `firm_data`: firm-level answers from the firm profile (see
    firm_profile.schedule_fill_data) -- ABN, insurances, certifications,
    registered address. Absent or empty leaves every one of those labels
    placeholdered exactly as before."""

    def _s(key):
        return (state.get(key) or "").strip() if isinstance(state.get(key), str) else ""

    data = {
        "company_name": _s("bidder_name"),
        "project_name": _s("project_name"),
        "client_name": _s("client_name"),
        "tender_number": _s("tender_name"),
        "submission_date": _s("submission_date_input"),
        "contact_name": _s("letter_sender_name"),
        "contact_title": _s("letter_sender_title"),
        "contact_phone": _s("letter_sender_phone"),
        "contact_email": _s("letter_sender_email"),
        "contact_address": _s("letter_sender_address"),
    }

    # Firm-level facts, where the account has entered them. Merged rather
    # than overwritten: a value typed for THIS bid wins over the standing
    # profile, since someone who overrode the address on this project meant
    # it.
    for key, value in (firm_data or {}).items():
        if value and not data.get(key):
            data[key] = value
    if not data.get("contact_address") and data.get("company_address"):
        data["contact_address"] = data["company_address"]
    if not data.get("company_name") and data.get("company_legal_name"):
        data["company_name"] = data["company_legal_name"]

    # Personnel: the resourcing plan is the source of truth for the whole
    # person -- who they are (person_name), what they're doing here
    # (slot/custom_title) AND their credentials.
    #
    # This used to read qualifications out of state["team_members"], which no
    # code in the app has ever written to (it is initialised to [] in
    # 10_state_helpers.py and only ever read, as a name pool). The lookup
    # therefore always missed, and the Qualifications and Years columns of
    # every council personnel schedule came out as placeholders even for a
    # user who had carefully filled both in on the Team & Resourcing tab.
    # Those fields are ResourceAssignment.qualification / .rpeq_status /
    # .years_experience, all user-entered and never AI-guessed, so they are
    # safe to put straight into a legal form.
    personnel = []
    for assignment in (state.get("resource_plan") or []):
        name = (getattr(assignment, "person_name", "") or "").strip()
        if not name:
            continue
        role = (getattr(assignment, "custom_title", "") or "").strip() or (getattr(assignment, "slot", "") or "").strip()
        # Schedules give one "Qualifications" column and firms expect the
        # registration in it alongside the degree, so the two entered values
        # are joined -- but only the parts that actually exist, so a person
        # with a degree and no RPEQ doesn't get a trailing comma.
        quals = ", ".join(
            part for part in (
                (getattr(assignment, "qualification", "") or "").strip(),
                (getattr(assignment, "rpeq_status", "") or "").strip(),
            ) if part
        )
        personnel.append({
            "name": name,
            "role": role,
            "quals": quals,
            "years": (getattr(assignment, "years_experience", "") or "").strip(),
        })
    data["personnel"] = personnel

    references = []
    for ref in (state.get("reference_projects") or []):
        title = (getattr(ref, "title", "") or "").strip()
        if title:
            references.append({
                "title": title,
                "client": (getattr(ref, "client", "") or "").strip(),
                "description": (getattr(ref, "description", "") or "").strip(),
            })
    data["references"] = references

    fee_lines = []
    for line in (state.get("discipline_fee_lines") or []):
        discipline = (getattr(line, "discipline", "") or "").strip()
        if discipline:
            fee_lines.append({
                "discipline": discipline,
                "hours": getattr(line, "total_hours", 0.0) or 0.0,
                "rate": getattr(line, "rate_per_hour", 0.0) or 0.0,
                "amount": (getattr(line, "total_hours", 0.0) or 0.0) * (getattr(line, "rate_per_hour", 0.0) or 0.0),
            })
    data["fee_lines"] = fee_lines
    return data


# ---------------------------------------------------------------------------
# Label matching
# ---------------------------------------------------------------------------

def _normalise_label(label: str) -> str:
    label = (label or "").lower()
    label = re.sub(r"[*_:.]+", " ", label)
    label = re.sub(r"\(.*?\)", " ", label)   # "(if applicable)" etc.
    return re.sub(r"\s+", " ", label).strip()


def _is_never_known(label_norm: str) -> bool:
    return any(re.search(p, label_norm) for p in NEVER_KNOWN_PATTERNS)


def match_label(label: str, fill_data: dict) -> tuple[str | None, str | None]:
    """Returns (field_key, value). field_key None = unrecognised label.
    value None = recognised (or never-known) but no real value -> placeholder."""
    label_norm = _normalise_label(label)
    if not label_norm or len(label_norm) > 120:
        return None, None
    # Firm-level labels first: an ABN or an insurer IS knowable now, if the
    # firm profile holds it. Falls through to never_known when it doesn't,
    # which is the original behaviour.
    for field_key, synonyms in FIRM_FIELD_SYNONYMS:
        for syn in synonyms:
            if re.search(rf"(?<![a-z]){re.escape(syn)}(?:s|es)?(?![a-z])", label_norm):
                value = (fill_data.get(field_key) or "").strip()
                if value:
                    return field_key, value
                return "never_known", None
    if _is_never_known(label_norm):
        return "never_known", None
    for field_key, synonyms in FIELD_SYNONYMS:
        for syn in synonyms:
            if re.search(rf"(?<![a-z]){re.escape(syn)}(?:s|es)?(?![a-z])", label_norm):
                value = (fill_data.get(field_key) or "").strip()
                return field_key, (value or None)
    return None, None


def _looks_like_label(text: str) -> bool:
    """A short piece of text that reads like a form label, not content."""
    text = (text or "").strip()
    if not text or len(text) > 120:
        return False
    if text.endswith(":"):
        return True
    words = text.split()
    return 1 <= len(words) <= 12


# ---------------------------------------------------------------------------
# Detection
# ---------------------------------------------------------------------------

def looks_like_response_form(filename: str, file_bytes: bytes) -> bool:
    """Heuristic used to OFFER filling for a directly-uploaded file. ZIP
    intake already classifies schedules; this is for loose uploads."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext in ("xlsx", "xlsm"):
        return True
    if ext != "docx":
        return False
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        total, empty = 0, 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    total += 1
                    if not cell.text.strip():
                        empty += 1
        return total >= 6 and empty / max(total, 1) >= 0.3
    except Exception:
        return False


# ---------------------------------------------------------------------------
# DOCX filling
# ---------------------------------------------------------------------------

def _write_into_cell(cell, value: str) -> None:
    """Adds text to a (near-)empty cell while keeping the client's cell
    formatting: reuse the cell's first paragraph (its style, alignment and
    any run formatting template survive) rather than cell.text=, which
    replaces the whole paragraph tree."""
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = value
        for extra in paragraph.runs[1:]:
            extra.text = ""
    else:
        paragraph.add_run(value)


def _cell_texts(row) -> list[str]:
    return [c.text.strip() for c in row.cells]


def _fill_label_value_row(row, fill_data: dict, where: str, result: FillResult) -> None:
    """A row whose first cell is a label and whose remaining cell(s) are
    empty -> fill the cell right after the label."""
    texts = _cell_texts(row)
    label = texts[0]
    if not label or not _looks_like_label(label):
        return
    # All non-first cells must be empty (a filled row is client content, and
    # this tool NEVER overwrites what's already there -- user edits included).
    target_index = None
    for i in range(1, len(row.cells)):
        if texts[i]:
            return
        if target_index is None:
            target_index = i
    if target_index is None:
        return
    field_key, value = match_label(label, fill_data)
    if value:
        _write_into_cell(row.cells[target_index], value)
        result.filled.append({"where": where, "label": label, "value": value})
    else:
        placeholder = make_placeholder(label)
        _write_into_cell(row.cells[target_index], placeholder)
        result.placeholdered.append({"where": where, "label": label})


def _map_header_columns(header_texts: list[str], column_synonyms: dict) -> dict[int, str]:
    """Match a table's header row against a column-synonym map. Returns
    {column_index: attribute}."""
    mapping = {}
    for i, text in enumerate(header_texts):
        norm = _normalise_label(text)
        if not norm:
            continue
        for attr, synonyms in column_synonyms.items():
            if any(re.search(rf"(?<![a-z]){re.escape(s)}(?:s|es)?(?![a-z])", norm) for s in synonyms):
                if attr not in mapping.values():
                    mapping[i] = attr
                break
    return mapping


def _fill_roster_table(table, items: list[dict], column_synonyms: dict,
                       where: str, result: FillResult, label_for_empty: str) -> bool:
    """Fills a header-row + empty-body table (key personnel, reference
    projects) from a list of dicts. Returns True if this table matched."""
    if len(table.rows) < 2:
        return False
    header = _cell_texts(table.rows[0])
    mapping = _map_header_columns(header, column_synonyms)
    # Demand at least 2 recognised columns including the primary one, and an
    # empty body -- otherwise this is a content table we must not touch.
    primary = "name" if "name" in column_synonyms else "title"
    if len(mapping) < 2 or primary not in mapping.values():
        return False
    body_rows = table.rows[1:]
    if any(any(_cell_texts(r)) for r in body_rows):
        return False  # body already has content -- client's or user's; hands off

    for row_index, row in enumerate(body_rows):
        if row_index < len(items):
            item = items[row_index]
            for col, attr in mapping.items():
                value = (item.get(attr) or "").strip()
                if value:
                    _write_into_cell(row.cells[col], value)
                    result.filled.append({"where": f"{where} row {row_index + 2}",
                                          "label": header[col] or attr, "value": value})
                else:
                    _write_into_cell(row.cells[col], make_placeholder(header[col] or attr))
                    result.placeholdered.append({"where": f"{where} row {row_index + 2}",
                                                 "label": header[col] or attr})
        elif row_index == len(items):
            # First spare row: one clear placeholder, not a wall of them.
            _write_into_cell(row.cells[list(mapping.keys())[0]],
                             make_placeholder(f"further {label_for_empty} if required"))
            result.placeholdered.append({"where": f"{where} row {row_index + 2}",
                                         "label": f"further {label_for_empty}"})
            break
    if not items and body_rows:
        # No data at all -- placeholder the first row's recognised columns so
        # the schedule visibly says what's missing.
        row = body_rows[0]
        for col, attr in mapping.items():
            _write_into_cell(row.cells[col], make_placeholder(header[col] or attr))
            result.placeholdered.append({"where": f"{where} row 2", "label": header[col] or attr})
    return True


_INLINE_BLANK = re.compile(r"^(?P<label>[^_.]{3,100}?):?\s*(?P<blank>_{3,}|\.{4,})\s*$")


def _fill_paragraph_blanks(paragraph, fill_data: dict, where: str, result: FillResult) -> None:
    """Fills 'Label: ______' / 'Label: ......' lines in body paragraphs,
    preserving the paragraph by editing run text only."""
    match = _INLINE_BLANK.match(paragraph.text.strip())
    if not match:
        return
    label = match.group("label").strip()
    if not _looks_like_label(label + ":"):
        return
    field_key, value = match_label(label, fill_data)
    new_tail = value if value else make_placeholder(label)
    # Replace the run(s) containing the blank; keep the label's runs intact.
    blank_pattern = re.compile(r"_{3,}|\.{4,}")
    replaced = False
    for run in paragraph.runs:
        if blank_pattern.search(run.text):
            if not replaced:
                new_text = blank_pattern.sub(new_tail, run.text, count=1)
                # The blank usually follows "Label: " -- avoid ending up
                # with a double space once the underscores become text.
                run.text = re.sub(r"  +", " ", new_text)
            else:
                run.text = blank_pattern.sub("", run.text)
            replaced = True
    if replaced:
        if value:
            result.filled.append({"where": where, "label": label, "value": value})
        else:
            result.placeholdered.append({"where": where, "label": label})


def fill_docx_schedule(filename: str, file_bytes: bytes, fill_data: dict) -> FillResult:
    result = FillResult(filename=filename, file_bytes=None, kind="docx")
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        result.error = (
            f"'{filename}' couldn't be opened as a Word document -- it may be corrupted or an "
            f"older .doc file renamed to .docx. ({str(exc)[:100]})"
        )
        return result

    try:
        for t_index, table in enumerate(doc.tables, start=1):
            where = f"table {t_index}"
            # Roster-style tables first (they'd otherwise look like many
            # label rows); personnel, then reference projects.
            if _fill_roster_table(table, fill_data.get("personnel") or [], ROSTER_COLUMNS,
                                  where, result, "personnel"):
                continue
            if _fill_roster_table(table, fill_data.get("references") or [], REFERENCE_COLUMNS,
                                  where, result, "reference projects"):
                continue
            for row in table.rows:
                _fill_label_value_row(row, fill_data, where, result)

        for p_index, paragraph in enumerate(doc.paragraphs, start=1):
            _fill_paragraph_blanks(paragraph, fill_data, f"paragraph {p_index}", result)

        out = io.BytesIO()
        doc.save(out)
        result.file_bytes = out.getvalue()
    except Exception as exc:
        result.error = (
            f"Something went wrong while filling '{filename}' -- the original file is untouched. "
            f"({str(exc)[:120]}) Email it to hello@civilproposals.com and we'll take a look."
        )
        result.file_bytes = None
    return result


# ---------------------------------------------------------------------------
# XLSX filling
# ---------------------------------------------------------------------------

XLSX_MAX_SCAN_ROWS = 500
XLSX_MAX_SCAN_COLS = 30
XLSX_MAX_PLACEHOLDERS_PER_SHEET = 60


def fill_xlsx_schedule(filename: str, file_bytes: bytes, fill_data: dict) -> FillResult:
    result = FillResult(filename=filename, file_bytes=None, kind="xlsx")
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes))  # values+formatting preserved on save
    except Exception as exc:
        result.error = (
            f"'{filename}' couldn't be opened as an Excel workbook -- it may be corrupted, "
            f"password-protected, or an older .xls file renamed. ({str(exc)[:100]}) "
            f"Email it to hello@civilproposals.com if it keeps failing."
        )
        return result

    try:
        for ws in wb.worksheets:
            placeholders_here = 0
            max_row = min(ws.max_row or 0, XLSX_MAX_SCAN_ROWS)
            max_col = min(ws.max_column or 0, XLSX_MAX_SCAN_COLS)
            for row in ws.iter_rows(min_row=1, max_row=max_row, max_col=max_col):
                for cell in row:
                    label = cell.value if isinstance(cell.value, str) else None
                    if not label or not _looks_like_label(label):
                        continue
                    target = ws.cell(row=cell.row, column=cell.column + 1)
                    if target.value not in (None, ""):
                        continue  # occupied -- never overwrite
                    if isinstance(target, openpyxl.cell.cell.MergedCell):
                        continue
                    field_key, value = match_label(label, fill_data)
                    where = f"{ws.title}!{target.coordinate}"
                    if value:
                        target.value = value
                        result.filled.append({"where": where, "label": label, "value": value})
                    elif field_key is not None or label.strip().endswith(":"):
                        # Recognised-but-unknown, never-known, or an explicit
                        # "Label:" cell -- placeholder it. Unrecognised prose
                        # cells are left alone (a spreadsheet is full of text
                        # that isn't a form field).
                        if placeholders_here < XLSX_MAX_PLACEHOLDERS_PER_SHEET:
                            target.value = make_placeholder(label)
                            result.placeholdered.append({"where": where, "label": label})
                            placeholders_here += 1

        out = io.BytesIO()
        wb.save(out)
        result.file_bytes = out.getvalue()
    except Exception as exc:
        result.error = (
            f"Something went wrong while filling '{filename}' -- the original file is untouched. "
            f"({str(exc)[:120]}) Email it to hello@civilproposals.com and we'll take a look."
        )
        result.file_bytes = None
    return result


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

def fill_schedule(filename: str, file_bytes: bytes, fill_data: dict) -> FillResult:
    """Fill one returnable schedule. Never raises."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    try:
        if ext == "docx":
            return fill_docx_schedule(filename, file_bytes, fill_data)
        if ext in ("xlsx", "xlsm"):
            return fill_xlsx_schedule(filename, file_bytes, fill_data)
        result = FillResult(filename=filename, file_bytes=None, kind="unsupported")
        result.error = (
            f"'.{ext}' schedules can't be filled automatically yet -- only .docx and "
            f".xlsx are supported so far. (.doc/.xls: re-save as .docx/.xlsx in Office "
            f"first; PDFs: email them to hello@civilproposals.com.)"
        )
        return result
    except Exception as exc:  # absolute backstop -- see module docstring
        result = FillResult(filename=filename, file_bytes=None, kind="unsupported")
        result.error = f"Unexpected problem filling '{filename}': {str(exc)[:120]}. The original file is untouched."
        return result


def filled_filename(filename: str) -> str:
    base, dot, ext = filename.rpartition(".")
    return f"{base} -- FILLED.{ext}" if dot else f"{filename} -- FILLED"
