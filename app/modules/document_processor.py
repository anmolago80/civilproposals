"""
document_processor.py

Extracts and prepares text from uploaded tender/company documents.

Supported inputs: PDF, DOCX, TXT.

Design notes
------------
- PyMuPDF (fitz) is used as the primary PDF engine because it can also read
  existing annotations (highlights, sticky notes, freetext callouts) -- this
  is what powers the "mark up the brief yourself, upload it, and the tool
  reads your comments" workflow.
- pdfplumber is used as a secondary pass purely for table extraction, since
  it tends to produce cleaner table grids than PyMuPDF for this purpose.
- Nothing here calls the AI. This module's only job is turning uploaded
  files into clean text + structured extras (headings, tables, annotations,
  approximate page references) for tender_analyser.py to work with.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field


# ---------------------------------------------------------------------------
# Data shapes
# ---------------------------------------------------------------------------

@dataclass
class ExtractedDocument:
    """Container for everything pulled out of an uploaded document."""
    filename: str
    text: str
    page_count: int = 0
    headings: list[str] = field(default_factory=list)
    tables: list[dict] = field(default_factory=list)   # {"page": int, "rows": [[...]]}
    page_texts: list[str] = field(default_factory=list)  # index 0 == page 1
    annotations: list[dict] = field(default_factory=list)  # see extract_annotations_from_pdf
    warning: str | None = None
    # True when any of this document's text came from OCR of scanned pages
    # rather than the PDF's own text layer -- see _ocr_scanned_pages(). The
    # UI and the exported documents both surface an "OCR-derived -- verify
    # carefully" notice off the back of this flag, because OCR can misread
    # numbers, dates, and names. ocr_pages is the 1-indexed pages affected.
    ocr_used: bool = False
    ocr_pages: list[int] = field(default_factory=list)


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------

# Above this many pages, the heading/table scan (font-metric heading detection +
# a page-by-page pdfplumber table pass) is too slow to run inline in the app --
# on a large PDF it can take minutes and looks like a freeze. Past this limit we
# extract the full text (fast, via PyMuPDF) and skip the structure scan.
STRUCTURE_SCAN_PAGE_LIMIT = 40

# --- Scanned-PDF OCR fallback ---------------------------------------------
# A page whose extracted text (whitespace-stripped) is shorter than this is
# treated as effectively image-only: real brief pages of body text are
# thousands of characters, while a scanned page's text layer is empty or
# just a few stray artefact characters. Deliberately low so a genuinely
# sparse-but-real page (a divider with one heading) isn't needlessly OCR'd.
OCR_MIN_CHARS_PER_PAGE = 40

# Rendering DPI for OCR -- 200 is Tesseract's sweet spot for body text
# (higher costs seconds per page for little accuracy gain on A4 briefs).
OCR_RENDER_DPI = 200

# Hard ceiling on how many pages get OCR'd inline, so a 500-page scanned
# appendix can't freeze the app for half an hour. Pages past the cap are
# reported in the warning rather than silently skipped.
OCR_MAX_PAGES = 60

# The standing tag wording -- referenced by app.py (UI badge) and
# export_docx.py (in-document notice) so all three stay in sync.
OCR_VERIFY_TAG = "OCR-derived -- verify carefully"

OCR_UNAVAILABLE_MESSAGE = (
    "This brief looks scanned -- its pages are images rather than selectable text, and "
    "text recognition (OCR) isn't available on the server right now, so nothing could be "
    "read from it. Try re-exporting the PDF with selectable text (most scanners and PDF "
    "tools have an 'OCR' or 'searchable PDF' option), or email the file to "
    "hello@civilproposals.com and we'll process it for you."
)


def ocr_available() -> bool:
    """True when pytesseract AND the tesseract binary are both usable.
    Cached after the first real check -- the answer can't change within one
    process lifetime, and the binary check shells out."""
    global _OCR_AVAILABLE
    if _OCR_AVAILABLE is None:
        try:
            import pytesseract
            pytesseract.get_tesseract_version()
            _OCR_AVAILABLE = True
        except Exception:
            _OCR_AVAILABLE = False
    return _OCR_AVAILABLE


_OCR_AVAILABLE: bool | None = None


def _ocr_pdf_page(page) -> str:
    """OCR one PyMuPDF page by rendering it to an image first. Returns ""
    on any failure -- the caller treats that the same as 'no text found'."""
    try:
        import pytesseract
        from PIL import Image

        pix = page.get_pixmap(dpi=OCR_RENDER_DPI)
        image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        return (pytesseract.image_to_string(image) or "").strip()
    except Exception:
        return ""


def extract_text_from_pdf(file_bytes: bytes, filename: str = "document.pdf",
                          include_structure: bool = True) -> ExtractedDocument:
    """
    Extract text from a PDF, plus (optionally) approximate headings and tables.

    `include_structure=False` returns text only -- fast -- for cases where the
    caller only needs the words (e.g. company reference material). Even when
    True, the slow heading/table scan is automatically skipped for PDFs larger
    than STRUCTURE_SCAN_PAGE_LIMIT pages so a big document can't freeze the app;
    the full text is always extracted regardless.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    page_count = doc.page_count
    do_structure = include_structure and page_count <= STRUCTURE_SCAN_PAGE_LIMIT

    page_texts: list[str] = []
    headings: list[str] = []
    for page in doc:
        page_texts.append(page.get_text("text") or "")
        if do_structure:
            headings.extend(_guess_headings_from_page(page))

    # --- Scanned-page detection + OCR fallback ---------------------------
    # A page with (almost) no text layer is effectively an image: a scan, a
    # photographed brief, or an image-only export. Without this, such a
    # brief used to sail through "successfully" with an empty analysis.
    low_text_pages = [
        i for i, t in enumerate(page_texts) if len((t or "").strip()) < OCR_MIN_CHARS_PER_PAGE
    ]
    ocr_pages: list[int] = []
    ocr_skipped_pages: list[int] = []
    ocr_needed = bool(low_text_pages)
    whole_doc_unreadable = len(low_text_pages) == page_count and page_count > 0

    if ocr_needed and ocr_available():
        for n, page_index in enumerate(low_text_pages):
            if n >= OCR_MAX_PAGES:
                ocr_skipped_pages.append(page_index + 1)
                continue
            ocr_text = _ocr_pdf_page(doc[page_index])
            if ocr_text:
                page_texts[page_index] = ocr_text
                ocr_pages.append(page_index + 1)

    full_text = "\n\n".join(
        f"[Page {i + 1}]\n{t}" for i, t in enumerate(page_texts)
    )
    doc.close()

    tables = _extract_tables_with_pdfplumber(file_bytes) if do_structure else []
    warnings: list[str] = []
    if include_structure and not do_structure:
        warnings.append(
            f"Large PDF ({page_count} pages): skipped the slower heading/table scan to keep "
            f"things responsive. The full text was still extracted and is used for analysis."
        )

    if ocr_needed and not ocr_available():
        if whole_doc_unreadable:
            # Nothing usable at all -- return empty text plus the plain-
            # language message, so the caller's "warning and not text" path
            # shows a clear error instead of proceeding to an empty analysis.
            return ExtractedDocument(
                filename=filename, text="", page_count=page_count,
                warning=OCR_UNAVAILABLE_MESSAGE,
            )
        warnings.append(
            f"{len(low_text_pages)} of this PDF's {page_count} pages look scanned (images "
            f"rather than selectable text) and text recognition (OCR) isn't available on the "
            f"server right now, so those pages could not be read: "
            f"page(s) {_format_page_list([p + 1 for p in low_text_pages])}. The other pages "
            f"were extracted normally. If those scanned pages matter, re-export the PDF with "
            f"selectable text or email it to hello@civilproposals.com."
        )
    elif ocr_pages:
        warnings.append(
            f"Text recognition (OCR) was used to read {len(ocr_pages)} scanned page(s) "
            f"(page {_format_page_list(ocr_pages)}). {OCR_VERIFY_TAG}: OCR can misread "
            f"numbers, dates, and names, so double-check anything important against the "
            f"original document."
        )
        if ocr_skipped_pages:
            warnings.append(
                f"{len(ocr_skipped_pages)} further scanned page(s) were NOT read "
                f"(page {_format_page_list(ocr_skipped_pages)}) -- OCR is capped at "
                f"{OCR_MAX_PAGES} pages per document to keep the app responsive."
            )
    elif ocr_needed and whole_doc_unreadable:
        # OCR is available but produced nothing on any page (e.g. blank
        # scans, or an image format Tesseract couldn't handle) -- same
        # honest dead-end as OCR being unavailable, never an empty analysis.
        return ExtractedDocument(
            filename=filename, text="", page_count=page_count,
            warning=(
                "This brief looks scanned, and text recognition (OCR) couldn't read anything "
                "usable from its pages. Try re-exporting the PDF with selectable text, or "
                "email the file to hello@civilproposals.com and we'll process it for you."
            ),
        )

    return ExtractedDocument(
        filename=filename,
        text=full_text,
        page_count=page_count,
        headings=_dedupe_preserve_order(headings)[:200],
        tables=tables,
        page_texts=page_texts,
        warning="\n\n".join(warnings) if warnings else None,
        ocr_used=bool(ocr_pages),
        ocr_pages=ocr_pages,
    )


def _format_page_list(pages: list[int], max_shown: int = 12) -> str:
    pages = sorted(pages)
    shown = ", ".join(str(p) for p in pages[:max_shown])
    if len(pages) > max_shown:
        shown += f" and {len(pages) - max_shown} more"
    return shown


def _guess_headings_from_page(page) -> list[str]:
    """Very rough heading detector: short lines rendered in a larger-than-body font."""
    found = []
    try:
        info = page.get_text("dict")
        body_sizes = []
        for block in info.get("blocks", []):
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    body_sizes.append(span.get("size", 0))
        if not body_sizes:
            return found
        median_size = sorted(body_sizes)[len(body_sizes) // 2]
        for block in info.get("blocks", []):
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                if not spans:
                    continue
                line_text = "".join(s.get("text", "") for s in spans).strip()
                max_size = max(s.get("size", 0) for s in spans)
                is_bold = any("Bold" in s.get("font", "") for s in spans)
                if (
                    line_text
                    and 3 <= len(line_text) <= 120
                    and (max_size >= median_size * 1.15 or is_bold)
                    and not line_text.endswith((".", ",", ";"))
                ):
                    found.append(line_text)
    except Exception:
        pass
    return found


def _extract_tables_with_pdfplumber(file_bytes: bytes) -> list[dict]:
    """Best-effort table extraction. Returns [] rather than raising on failure."""
    tables_out = []
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                try:
                    page_tables = page.extract_tables()
                except Exception:
                    page_tables = []
                for rows in page_tables or []:
                    cleaned_rows = [
                        [(cell or "").strip() for cell in row] for row in rows
                    ]
                    if any(any(cell for cell in row) for row in cleaned_rows):
                        tables_out.append({"page": page_number, "rows": cleaned_rows})
    except Exception:
        # Table extraction is a bonus, not a requirement -- never fail the whole upload over it.
        pass
    return tables_out


def extract_annotations_from_pdf(file_bytes: bytes, filename: str = "document.pdf") -> list[dict]:
    """
    Pull existing highlights / sticky notes / freetext callouts out of a PDF the user has
    already marked up (e.g. in Acrobat or Word) while reading the brief.

    Returns a list of dicts:
        {
            "page": int,               # 1-indexed
            "type": "Highlight" | "Text" | "FreeText" | "Squiggly" | ...,
            "comment": str,             # the note/comment text, if any
            "highlighted_text": str,    # the underlying text the annotation covers, if any
        }
    """
    import fitz  # PyMuPDF

    results: list[dict] = []
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page_index, page in enumerate(doc):
            for annot in page.annots() or []:
                info = annot.info or {}
                comment = (info.get("content") or "").strip()
                annot_type = info.get("subject") or annot.type[1]

                highlighted_text = ""
                try:
                    vertices = annot.vertices
                    if vertices:
                        # Highlight/underline/squiggly annotations store quad points;
                        # group them in fours and pull the text under each quad.
                        quad_points = vertices
                        for i in range(0, len(quad_points), 4):
                            quad = quad_points[i:i + 4]
                            if len(quad) == 4:
                                rect = fitz.Quad(quad).rect
                                highlighted_text += page.get_textbox(rect) + " "
                    elif annot.rect:
                        highlighted_text = page.get_textbox(annot.rect)
                except Exception:
                    highlighted_text = ""

                highlighted_text = highlighted_text.strip()
                if not comment and not highlighted_text:
                    continue  # skip empty/decorative annotations

                results.append({
                    "page": page_index + 1,
                    "type": annot_type,
                    "comment": comment,
                    "highlighted_text": highlighted_text,
                })
        doc.close()
    except Exception:
        pass
    return results


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def extract_text_from_docx(file_bytes: bytes, filename: str = "document.docx") -> ExtractedDocument:
    """Extract text, headings, and tables from a DOCX file.

    Also pulls text from places python-docx's Document.paragraphs never visits:
    text boxes/shapes (<w:txbxContent>, nested inside a run's drawing and
    invisible to normal paragraph traversal), content controls (<w:sdt> --
    Word's "structured document tag" form fields, common in template-driven CV
    documents for things like a name banner or a "proposed role" field), and
    section headers/footers. This matters a lot for CV templates in
    particular -- it's extremely common for the person's name to sit in a
    coloured banner built as a text box, a content-control field, or a
    repeated page header, with the body text starting straight into
    qualifications/experience. Without this, a CV extracts "cleanly" (no
    errors, plausible-looking text) while silently missing the one thing
    everything downstream (name matching, resourcing dropdowns, drafted
    profiles) needs most: the person's actual name.

    Table content is also folded into the plain text (not just kept in the
    separate `tables` structure below). Several real-world CV templates --
    including the TMR/DTMR "Curriculum vitae" template -- put essentially all
    of a person's actual content (qualifications, RPEQ/registration numbers,
    stated years of experience, career summary, project experience) inside
    Word tables rather than body paragraphs. Without this, the extracted text
    for those CVs is close to an empty template shell -- just section-heading
    labels -- and every downstream step (name matching, drafted profiles,
    "years of experience") silently has almost nothing real to work with, even
    though the CV plainly states the fact when read by a person."""
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn

    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs_text = []
    headings = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        paragraphs_text.append(text)
        style_name = (para.style.name if para.style else "") or ""
        if style_name.lower().startswith("heading") or style_name.lower() == "title":
            headings.append(text)

    # Text boxes/shapes and content-control (<w:sdt>) fields anywhere in the
    # body -- put first, since a name banner is almost always the visual top
    # of the document even though python-docx's plain paragraph order may not
    # reflect that.
    textbox_texts = _extract_textbox_paragraphs(doc.element.body, qn)
    sdt_texts = _extract_sdt_paragraphs(doc.element.body, qn)

    # Section headers/footers (e.g. a name banner repeated on every page).
    header_footer_texts: list[str] = []
    for section in doc.sections:
        for part in (section.header, section.footer):
            for para in part.paragraphs:
                t = para.text.strip()
                if t:
                    header_footer_texts.append(t)
            header_footer_texts += _extract_textbox_paragraphs(part._element, qn)
            header_footer_texts += _extract_sdt_paragraphs(part._element, qn)
    header_footer_texts = _dedupe_preserve_order(header_footer_texts)

    banner_texts = _dedupe_preserve_order(textbox_texts + sdt_texts)
    paragraphs_text = banner_texts + header_footer_texts + paragraphs_text

    tables = []
    table_texts: list[str] = []
    for t_index, table in enumerate(doc.tables):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if rows:
            tables.append({"page": None, "rows": rows})
            table_text = _table_rows_to_text(rows)
            if table_text:
                table_texts.append(table_text)

    # Fold table content into the plain text too, after the paragraph/banner
    # text -- see the docstring above for why this matters. Rendered as one
    # block per table rather than interleaved at the exact original position
    # (python-docx doesn't expose a cheap way to interleave body paragraphs
    # and tables in true document order); every fact still ends up in the
    # text, just not always in its original visual position.
    full_text = "\n".join(paragraphs_text + table_texts)

    return ExtractedDocument(
        filename=filename,
        text=full_text,
        page_count=0,  # DOCX has no reliable page count without rendering
        headings=_dedupe_preserve_order(headings)[:200],
        tables=tables,
        page_texts=[full_text],  # DOCX is treated as one "page" of text
    )


def _extract_textbox_paragraphs(element, qn) -> list[str]:
    """Every paragraph of text inside <w:txbxContent> elements anywhere under
    `element` (a document body, or a header/footer element), in document order.
    This is the only reliable way to reach text-box/shape content with
    python-docx, which has no first-class API for it."""
    texts = []
    for txbx in element.iter(qn("w:txbxContent")):
        for para in txbx.iter(qn("w:p")):
            run_text = "".join((node.text or "") for node in para.iter(qn("w:t")))
            run_text = run_text.strip()
            if run_text:
                texts.append(run_text)
    return texts


def _extract_sdt_paragraphs(element, qn) -> list[str]:
    """Every paragraph of text inside <w:sdt> (structured document tag /
    content control) elements anywhere under `element`, in document order.

    Word content controls -- form-field-style elements used heavily by
    template-driven documents for things like a name banner, a "proposed
    role" field, or a date field -- wrap their run(s) in
    <w:sdt><w:sdtContent>... rather than as a direct child <w:r> of the
    paragraph. python-docx's ordinary paragraph/run traversal (doc.paragraphs
    / paragraph.text) only ever sees DIRECT child <w:r> elements of a <w:p>,
    so this text is completely invisible to it -- and unlike a missing
    qualification or project bullet, a missing NAME (which is exactly what
    tends to sit in one of these fields on a CV) silently breaks every
    downstream step that needs to match a person to their own CV file."""
    texts = []
    for sdt in element.iter(qn("w:sdt")):
        content = sdt.find(qn("w:sdtContent"))
        if content is None:
            continue
        paras = list(content.iter(qn("w:p")))
        if paras:
            # Block-level content control -- <w:sdtContent> wraps one or more
            # whole <w:p> paragraphs.
            for para in paras:
                run_text = "".join((node.text or "") for node in para.iter(qn("w:t")))
                run_text = run_text.strip()
                if run_text:
                    texts.append(run_text)
        else:
            # Inline/run-level content control -- <w:sdtContent> wraps a bare
            # <w:r> run (or several) with no <w:p> in between, which is the
            # more common case for a single-field CV banner like a name or
            # role. Grab all text directly under the content node as one unit.
            run_text = "".join((node.text or "") for node in content.iter(qn("w:t")))
            run_text = run_text.strip()
            if run_text:
                texts.append(run_text)
    return texts


def _table_rows_to_text(rows: list[list[str]]) -> str:
    """Flatten a table's rows into plain text, one line per row.

    python-docx reports a horizontally-merged cell as the SAME text repeated
    once per underlying grid column it spans (there's no first-class "merged
    cell" concept in the API) -- without handling that, a table like a CV's
    project-experience grid (common merged-cell layout) would repeat the same
    paragraph 4-6 times in the extracted text. Collapsing consecutive
    duplicate cells within a row fixes that while leaving genuinely distinct
    cells (including a cell that happens to equal a NON-adjacent cell
    elsewhere in the row) untouched."""
    lines = []
    for row in rows:
        deduped: list[str] = []
        for cell in row:
            cell = (cell or "").strip()
            if not cell:
                continue
            if deduped and deduped[-1] == cell:
                continue
            deduped.append(cell)
        if deduped:
            lines.append(" | ".join(deduped))
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# TXT extraction
# ---------------------------------------------------------------------------

def extract_text_from_txt(file_bytes: bytes, filename: str = "document.txt") -> ExtractedDocument:
    """Extract text from a plain text file, trying a couple of common encodings."""
    text = ""
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            text = file_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    return ExtractedDocument(filename=filename, text=text, page_texts=[text])


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {"pdf", "docx", "txt"}


def _read_upload(uploaded_file):
    """(filename, extension, file_bytes) from a Streamlit UploadedFile-like object."""
    filename = getattr(uploaded_file, "name", "document")
    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    try:
        file_bytes = uploaded_file.getvalue()
    except AttributeError:
        file_bytes = uploaded_file.read()
    return filename, extension, file_bytes


def _unsupported(filename: str, extension: str) -> ExtractedDocument:
    return ExtractedDocument(
        filename=filename, text="",
        warning=(f"Unsupported file type '.{extension}'. Supported formats are "
                 f"PDF, DOCX, and TXT. This file was skipped."),
    )


def extract_text_from_file(uploaded_file) -> ExtractedDocument:
    """
    Full extraction (text + headings + tables + PDF annotations). Used for the
    tender brief, where the mark-up/annotation and structure detail matter.

    `uploaded_file` is expected to be a Streamlit UploadedFile (has .name and
    .getvalue()), but anything with the same interface works.
    """
    filename, extension, file_bytes = _read_upload(uploaded_file)
    if extension not in SUPPORTED_EXTENSIONS:
        return _unsupported(filename, extension)
    try:
        if extension == "pdf":
            doc = extract_text_from_pdf(file_bytes, filename)
            doc.annotations = extract_annotations_from_pdf(file_bytes, filename)
            return doc
        elif extension == "docx":
            return extract_text_from_docx(file_bytes, filename)
        else:  # txt
            return extract_text_from_txt(file_bytes, filename)
    except Exception as exc:
        return ExtractedDocument(filename=filename, text="", warning=f"Could not read '{filename}': {exc}")


def extract_plain_text_from_file(uploaded_file) -> ExtractedDocument:
    """
    Fast, text-only extraction for company material (profile, previous proposals,
    references, CV library, boilerplate), where only the raw words are ever used --
    never headings, tables, or annotations. Skips the slow pdfplumber table scan
    and font-metric heading detection entirely, so a large company PDF (e.g. a
    100-page previous proposal) extracts in a second or two instead of freezing
    the app for minutes.
    """
    filename, extension, file_bytes = _read_upload(uploaded_file)
    if extension not in SUPPORTED_EXTENSIONS:
        return _unsupported(filename, extension)
    try:
        if extension == "pdf":
            return extract_text_from_pdf(file_bytes, filename, include_structure=False)
        elif extension == "docx":
            return extract_text_from_docx(file_bytes, filename)
        else:  # txt
            return extract_text_from_txt(file_bytes, filename)
    except Exception as exc:
        return ExtractedDocument(filename=filename, text="", warning=f"Could not read '{filename}': {exc}")


def combine_extracted_documents(docs: list[ExtractedDocument]) -> ExtractedDocument:
    """
    Merge several per-file ExtractedDocument objects (e.g. a main RFT plus
    addenda/schedules/annexures uploaded as separate files) into a single
    ExtractedDocument, so the rest of the app can keep treating "the tender
    brief" as one object.

    - text: each file's text is kept intact and separated by a clear
      "===== filename =====" header, so the AI can still tell where one
      document ends and the next begins.
    - annotations/tables: concatenated, with a "source_file" key added to
      each entry, since page numbers are only meaningful within their own
      file once multiple PDFs are combined.
    - headings/page_texts: concatenated as-is.
    - page_count: summed across files that report one (DOCX/TXT contribute 0).
    - warning: any per-file warnings are combined into one message so a
      problem with one file (e.g. an unsupported type) doesn't get silently
      dropped just because the other files extracted fine.

    If given a single document, returns it unchanged (no wrapping/renaming)
    so single-file uploads behave exactly as before.
    """
    docs = [d for d in docs if d is not None]
    if not docs:
        return ExtractedDocument(filename="(no files)", text="", warning="No files were provided.")
    if len(docs) == 1:
        return docs[0]

    filenames = [d.filename for d in docs]
    text_parts = []
    headings: list[str] = []
    tables: list[dict] = []
    page_texts: list[str] = []
    annotations: list[dict] = []
    warnings: list[str] = []
    page_count = 0

    ocr_used = False
    ocr_pages: list[int] = []
    for d in docs:
        if d.text:
            text_parts.append(f"===== {d.filename} =====\n\n{d.text}")
        headings.extend(d.headings)
        page_texts.extend(d.page_texts)
        page_count += d.page_count or 0
        for t in d.tables:
            tables.append({**t, "source_file": d.filename})
        for a in d.annotations:
            annotations.append({**a, "source_file": d.filename})
        if d.warning:
            warnings.append(f"{d.filename}: {d.warning}")
        if getattr(d, "ocr_used", False):
            ocr_used = True
            ocr_pages.extend(getattr(d, "ocr_pages", []) or [])

    return ExtractedDocument(
        filename=f"{len(docs)} files combined ({', '.join(filenames)})",
        text="\n\n\n".join(text_parts),
        page_count=page_count,
        headings=headings,
        tables=tables,
        page_texts=page_texts,
        annotations=annotations,
        warning="; ".join(warnings) if warnings else None,
        ocr_used=ocr_used,
        ocr_pages=ocr_pages,
    )


# ---------------------------------------------------------------------------
# Cleaning + chunking
# ---------------------------------------------------------------------------

def merge_extracted_material(existing: dict, updates: dict) -> dict:
    """
    Merge freshly-extracted per-file text into a "company material" category's
    existing per-file store (keyed by filename). Re-uploading a file replaces
    just that entry; every other file already in `existing` is left untouched.

    This matters because Streamlit's file_uploader only ever reports the set
    of files currently selected in the widget, not everything uploaded so far
    this session -- a naive "replace the whole category with what's currently
    selected" approach means uploading one or two corrected CVs silently wipes
    out everyone else's CV text. Keying by filename and merging instead means
    someone can re-upload just the files that were wrong (e.g. after a bug fix
    to text extraction) without having to reselect the entire library.
    """
    merged = dict(existing or {})
    merged.update(updates or {})
    return merged


def clean_extracted_text(text: str) -> str:
    """Normalise whitespace and strip common PDF extraction artefacts."""
    if not text:
        return ""
    text = text.replace(" ", " ")            # non-breaking spaces
    text = re.sub(r"[ \t]+", " ", text)            # collapse runs of spaces/tabs
    text = re.sub(r"\n{3,}", "\n\n", text)         # collapse excessive blank lines
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)   # rejoin hyphenated line-wraps
    return text.strip()


def split_text_into_chunks(text: str, chunk_size: int = 6000, overlap: int = 400) -> list[str]:
    """
    Split long text into overlapping chunks suitable for sending to an LLM.

    Splits on paragraph boundaries where possible so chunks don't cut a
    sentence in half; falls back to a hard character split if a single
    paragraph is itself longer than chunk_size.
    """
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]

    paragraphs = text.split("\n\n")
    chunks: list[str] = []
    current = ""

    for para in paragraphs:
        candidate = f"{current}\n\n{para}" if current else para
        if len(candidate) <= chunk_size:
            current = candidate
            continue

        if current:
            chunks.append(current)
        if len(para) > chunk_size:
            for i in range(0, len(para), chunk_size - overlap):
                chunks.append(para[i:i + chunk_size])
            current = ""
        else:
            current = para

    if current:
        chunks.append(current)

    # Add overlap between consecutive chunks for context continuity.
    overlapped = []
    for i, chunk in enumerate(chunks):
        if i == 0:
            overlapped.append(chunk)
        else:
            tail = chunks[i - 1][-overlap:] if overlap else ""
            overlapped.append(f"{tail}\n{chunk}")
    return overlapped


def _dedupe_preserve_order(items: list[str]) -> list[str]:
    seen = set()
    out = []
    for item in items:
        key = item.strip().lower()
        if key and key not in seen:
            seen.add(key)
            out.append(item.strip())
    return out
