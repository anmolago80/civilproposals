"""
export_docx.py

Assembles the final first-pass DOCX response pack as TWO separate documents,
via two entry points:

build_docx() -- the Proposal document itself, and only the proposal:
cover page -> TOC placeholder -> executive summary -> page allocation plan ->
indicative fee estimate by discipline (the %-based table) -> proposal
response sections (each starting on a new page, red guidance note first, then
first-pass draft, then graphic placeholders).

build_tender_summary_docx() -- a companion internal working document, kept
OUT of the proposal itself: title page -> TOC -> tender summary (brief
analysis) -> compliance matrix -> gap analysis -> review checklist -> user
input required list.

This is a first-pass PREPARATION pack, not submission-ready copy -- that's
stated on the cover page itself, not just in this docstring.
"""

from __future__ import annotations

import io
import re
import sys
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Emu

from modules import divider_designer
from modules import export_i18n

DEFAULT_FONT = "Arial"

RED = RGBColor(0xC0, 0x00, 0x00)
DARK_GREY = RGBColor(0x40, 0x40, 0x40)
RISK_COLORS = {"High": RGBColor(0xC0, 0x00, 0x00), "Medium": RGBColor(0xB8, 0x86, 0x00), "Low": RGBColor(0x2E, 0x7D, 0x32)}
RISK_SHADING = {"High": "F8D7DA", "Medium": "FFF3CD", "Low": "D4EDDA"}


def _theme_colours(proposal_theme: str | None) -> dict:
    """Same palette the divider/cover banners use (divider_designer.THEME_COLOURS), so
    headings, table headers, and accent lines match the generated graphics rather than
    defaulting to Word's generic blue theme."""
    palette = divider_designer.THEME_COLOURS.get(proposal_theme, divider_designer.THEME_COLOURS["Corporate"])
    primary = RGBColor(*palette["primary"])
    accent = RGBColor(*palette["accent"])
    # Minimalist's "primary" is a light background colour, not a text colour -- headings
    # need something dark enough to read, so fall back to a near-black for that theme only.
    heading_colour = RGBColor(0x2A, 0x2A, 0x2A) if proposal_theme == "Minimalist" else primary
    return {"primary": primary, "accent": accent, "heading": heading_colour}


def build_docx(
    project_info: dict,
    analysis,
    weighted_criteria: list,
    allocations: list,
    sections: list,
    guidance_notes: dict,
    drafts: dict,
    compliance_items: list,
    gap_items: list,
    graphics: list,
    fee_estimates: list | None = None,
    weighting_chart_png: bytes | None = None,
    cover_image_bytes: bytes | None = None,
    cover_theme_image_bytes: bytes | None = None,
    divider_images: dict[str, bytes] | None = None,
    resource_plan: list | None = None,
    org_chart_png: bytes | None = None,
    body_font: str | None = None,
    personnel_photos: dict[str, bytes] | None = None,
    reference_projects: list | None = None,
    reference_project_photos: dict[str, bytes] | None = None,
    discipline_fee_lines: list | None = None,
    executive_summary=None,
    fee_estimate_indicative_amounts: dict | None = None,
    team_intro=None,
    experience_intro=None,
    differentiator_text: str | None = None,
    sales_pitch_text: str | None = None,
    ocr_note: str | None = None,
    program_schedule: dict[str, list[bool]] | None = None,
    program_week_labels: list[str] | None = None,
    firm: dict | None = None,
    fee_sections_included: dict | None = None,
    scope_item_fees: list | None = None,
    output_language: str = "en",
) -> io.BytesIO:
    theme = _theme_colours(project_info.get("proposal_theme"))
    firm = firm or {}
    font = body_font or DEFAULT_FONT

    doc = Document()
    _set_base_styles(doc, theme, font)
    _add_page_numbers(doc)

    _build_cover_page(doc, project_info, cover_image_bytes, cover_theme_image_bytes, theme,
                      firm=firm, output_language=output_language)

    _add_ocr_notice(doc, ocr_note)

    _add_toc(doc, output_language)
    doc.add_page_break()

    _build_executive_summary(doc, executive_summary, project_info, theme,
                             differentiator_text=differentiator_text, output_language=output_language)
    doc.add_page_break()

    _build_page_allocation_plan(doc, sections, theme, output_language)
    doc.add_page_break()

    # Which fee presentations reach the proposal is the user's choice now
    # (Fee Estimate tab) rather than hardcoded per pack format.
    _fee_included = fee_sections(fee_sections_included)
    if fee_estimates and _fee_included["pct_split"]:
        _build_fee_estimate(doc, fee_estimates, analysis.fee_cap, theme, fee_estimate_indicative_amounts,
                            output_language=output_language)
        doc.add_page_break()

    _build_proposal_response(doc, sections, guidance_notes, drafts, graphics, divider_images,
                             analysis, resource_plan, org_chart_png, theme,
                             personnel_photos=personnel_photos,
                             reference_projects=reference_projects,
                             reference_project_photos=reference_project_photos,
                             discipline_fee_lines=discipline_fee_lines,
                             project_info=project_info,
                             team_intro=team_intro,
                             experience_intro=experience_intro,
                             sales_pitch_text=sales_pitch_text,
                             program_schedule=program_schedule,
                             program_week_labels=program_week_labels,
                             firm=firm,
                             fee_sections_included=_fee_included,
                             scope_item_fees=scope_item_fees,
                             output_language=output_language)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def build_tender_summary_docx(
    project_info: dict,
    analysis,
    weighting_chart_png: bytes | None,
    compliance_items: list,
    gap_items: list,
    sections: list,
    drafts: dict,
    body_font: str | None = None,
    ocr_note: str | None = None,
    document_placeholders: list | None = None,
    output_language: str = "en",
) -> io.BytesIO:
    """Builds the companion Tender Summary document -- everything about how the
    brief was read and how this response pack was put together, kept OUT of the
    main Proposal document so that one contains only the proposal itself: the
    tender summary (scope/objectives/criteria/deliverables/risks, exactly as
    extracted from the brief), the compliance matrix, the gap analysis, the
    review checklist, and the user-input-required list. This is an internal
    working document for the bid team, not something submitted to the client."""
    theme = _theme_colours(project_info.get("proposal_theme"))
    font = body_font or DEFAULT_FONT

    doc = Document()
    _set_base_styles(doc, theme, font)
    _add_page_numbers(doc)

    _build_tender_summary_title_page(doc, project_info, theme, output_language)
    doc.add_page_break()

    _add_ocr_notice(doc, ocr_note)

    _add_toc(doc, output_language)
    doc.add_page_break()

    _build_tender_summary(doc, analysis, weighting_chart_png, theme, output_language)
    doc.add_page_break()

    _build_compliance_matrix(doc, compliance_items, theme, output_language)
    doc.add_page_break()

    _build_gap_analysis(doc, gap_items, theme, output_language)
    doc.add_page_break()

    _build_review_checklist(doc, sections, output_language)
    doc.add_page_break()

    _build_user_input_list(doc, compliance_items, gap_items, drafts, document_placeholders,
                           output_language=output_language)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _build_tender_summary_title_page(doc: Document, project_info: dict, theme: dict,
                                     output_language: str = "en"):
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(export_i18n.export_t("heading_tender_summary", output_language))
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = theme["heading"]

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(project_info.get("project_name") or project_info.get("tender_name") or "")
    run.font.size = Pt(15)
    run.font.color.rgb = DARK_GREY

    if project_info.get("client_name"):
        client_p = doc.add_paragraph()
        client_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        client_p.add_run(project_info["client_name"]).font.size = Pt(12)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "Internal working document -- brief analysis, compliance matrix, gap analysis, and "
        "review checklist. Companion to the Proposal document, not part of what's submitted "
        f"to the client. Generated {datetime.now().strftime('%d %B %Y')}."
    )
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_GREY


def build_letter_docx(
    project_info: dict,
    sender: dict,
    analysis,
    understanding_text: str,
    methodology_text: str,
    resource_plan: list,
    personnel_photos: dict[str, bytes],
    program_schedule: dict[str, list[bool]],
    program_week_labels: list[str],
    terms_of_engagement_text: str,
    executive_summary=None,
    cover_image_bytes: bytes | None = None,
    cover_theme_image_bytes: bytes | None = None,
    fee_estimates: list | None = None,
    discipline_fee_lines: list | None = None,
    differentiator_text: str | None = None,
    sales_pitch_text: str | None = None,
    ocr_note: str | None = None,
    firm: dict | None = None,
    risk_register=None,
    fee_sections_included: dict | None = None,
    scope_item_fees: list | None = None,
    program_style: str | None = None,
    methodology_stages: list | None = None,
    program_start_date=None,
    output_language: str = "en",
) -> io.BytesIO:
    """
    Builds the Small Scope Proposal Response Pack -- the leaner, content-agnostic
    sibling of build_docx() for a small brief or an email-based fee request. Same
    underlying pipeline and no-invention discipline throughout, just a smaller
    shape: full-bleed cover page (reusing _build_cover_page, same as the Large
    Scope pack -- no section dividers here, this stays a short document), an
    Executive Summary as the warm opening, eight lean numbered sections, sign-off,
    review checklist. No ToC, compliance matrix, or per-section red guidance walls.

    There used to be a plain-text letterhead here instead (sender address block,
    Your ref/Our ref, date, recipient block, subject line, "Dear ..." salutation,
    "Thank you for the opportunity ..." opening) -- removed once the cover page
    and Executive Summary took over that job; a second warm opening paragraph
    right after the Executive Summary would've been redundant. sender is now only
    used for the closing "Regards" sign-off block.

    sender: {"name": str, "title": str, "phone": str, "email": str}
    resource_plan: list[resourcing.ResourceAssignment] -- the SAME Team & Resourcing
      plan the Large Scope pack's Key Personnel section reads, including any
      support members added under a discipline lead (see
      resourcing.letter_team_entries, which is what actually renders the Project
      Team section from this). personnel_photos: {person_name: headshot_bytes}
      -- also shared with the Large Scope pack, not a separate photo store.
    program_schedule: {scope_item_title: [bool per week]}
    fee_estimates: list[fee_estimation_engine.DisciplineFeeEstimate] | None -- the
      discipline fee % split from the Fees & Program tab's "Discipline fee split"
      section, if generated -- this is the fee split table that actually goes into
      the pack (the old per-scope-item fee table never did and has been dropped
      from here entirely). Only the percentage breakdown is included (never a $
      amount, unlike the Large Scope pack's Fee Estimate tab); omitted entirely
      from the export when not generated.
    discipline_fee_lines: list[resourcing.DisciplineFeeLine] | None -- the hours x
      rate discipline fee build-up from the Fees & Program tab, same model and same
      Fee summary table shape as the Large Scope pack's Commercial section
      (_build_commercial_section): a $ total per discipline only, never the
      underlying hours/rate. Omitted entirely from the export when not entered.
    """
    theme = _theme_colours(project_info.get("proposal_theme"))
    firm = firm or {}

    doc = Document()
    _set_base_styles(doc, theme)
    _add_page_numbers(doc)

    _build_cover_page(doc, project_info, cover_image_bytes, cover_theme_image_bytes, theme,
                      firm=firm, output_language=output_language)

    _add_ocr_notice(doc, ocr_note)
    _add_company_footer_line(doc.sections[-1].footer, project_info, firm, output_language)

    _build_executive_summary(doc, executive_summary, project_info, theme,
                             differentiator_text=differentiator_text, output_language=output_language)
    doc.add_page_break()

    doc.add_heading(export_i18n.export_t("heading_letter_intro", output_language), level=1)
    _add_letter_body_text(doc, understanding_text,
                          export_i18n.export_t("export_no_introduction", output_language), theme)

    doc.add_heading(export_i18n.export_t("heading_letter_scope_of_work", output_language), level=1)
    _build_letter_scope_of_work(doc, analysis.scope_items, output_language)

    doc.add_heading(export_i18n.export_t("heading_letter_methodology", output_language), level=1)
    _build_letter_methodology(doc, methodology_text, theme, sales_pitch_text=sales_pitch_text,
                              output_language=output_language)

    doc.add_heading(export_i18n.export_t("heading_letter_team", output_language), level=1)
    _build_letter_team(doc, resource_plan, personnel_photos, theme, output_language)

    doc.add_heading(export_i18n.export_t("heading_letter_fees", output_language), level=1)
    # Stable order, and only the presentations the user ticked on the Fee
    # Estimate tab. The defaults reproduce what this pack always exported.
    _letter_fee = fee_sections(fee_sections_included)
    if not any(_letter_fee.values()):
        _add_placeholder_paragraph(doc, export_i18n.export_t("export_fee_nothing_selected", output_language))
    if _letter_fee["pct_split"] and fee_estimates:
        _build_letter_fee_split(doc, fee_estimates, theme, output_language)
    if _letter_fee["discipline_buildup"] and discipline_fee_lines:
        _build_letter_fee_buildup(doc, discipline_fee_lines, theme, output_language)
    if _letter_fee["scope_buildup"]:
        _build_scope_item_fees(doc, scope_item_fees, theme, output_language=output_language)
    if any(_letter_fee.values()) and not discipline_fee_lines and not fee_estimates and not (
            _letter_fee["scope_buildup"] and scope_item_fees):
        # Both tables are optional, and with neither of them this heading
        # rendered with literally nothing underneath it -- a numbered "5.
        # Fees" section followed by "6. Program". A fee proposal that appears
        # to have deliberately said nothing about fees is worse than one that
        # visibly still needs them.
        _add_placeholder_paragraph(
            doc, export_i18n.export_t("export_no_fees_entered_letter", output_language),
        )

    doc.add_heading(export_i18n.export_t("heading_letter_program", output_language), level=1)
    _build_letter_program(
        doc, program_schedule, program_week_labels, theme,
        style=program_style, methodology_stages=methodology_stages,
        start_date=program_start_date, analysis=analysis,
        project_name=project_info.get("project_name", ""),
        client_name=project_info.get("client_name", ""),
        output_language=output_language,
    )

    doc.add_heading(export_i18n.export_t("heading_letter_assumptions", output_language), level=1)
    if analysis.assumptions:
        _add_bullets(doc, analysis.assumptions)
    else:
        _add_placeholder_paragraph(doc, export_i18n.export_t("export_no_assumptions", output_language))
    # The brief's risks used to go in here as raw bullets -- the client's own
    # sentences read back to the client who wrote them, saying nothing about
    # what they mean for delivery or what this firm would do about them. With
    # a reviewed register they become a real risk / impact / mitigation table.
    if getattr(risk_register, "entries", None):
        doc.add_heading(export_i18n.export_t("heading_risks_mitigation", output_language), level=2)
        _build_risk_table(doc, risk_register, theme)
    elif analysis.risks:
        doc.add_heading(export_i18n.export_t("heading_risks_noted_in_brief", output_language), level=2)
        _add_bullets(doc, analysis.risks)

    doc.add_heading(export_i18n.export_t("heading_letter_terms", output_language), level=1)
    _add_letter_body_text(doc, terms_of_engagement_text,
                          export_i18n.export_t("export_no_terms_of_engagement", output_language))

    _build_letter_signoff(doc, sender, output_language)

    doc.add_page_break()
    _build_letter_review_checklist(doc, output_language)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# Small Scope pack-specific sections
# ---------------------------------------------------------------------------

def _add_company_footer_line(footer, project_info: dict, firm: dict | None = None,
                             output_language: str = "en") -> None:
    """Second footer line, below the "Page X of Y" one _build_cover_page already
    wrote -- the bidder's registered company details, for the user to confirm
    before sending. Pre-fills the real bidder name where we have it; ABN and
    registered address are never invented, always a bracketed placeholder,
    same convention as every other "confirm before submission" note in this
    tool."""
    firm = firm or {}
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # With a firm profile this is the real "Firm | ABN ... | address" line on
    # every page. Without one it is the same red placeholder it always was --
    # and it stays red whenever ANY of the three parts is still missing, so a
    # half-filled profile can't produce a black line that looks finished.
    text = firm.get("footer_line")
    complete = bool(firm.get("footer_complete"))
    if not text:
        bidder = (project_info.get("bidder_name") or "").strip() or export_i18n.export_t(
            "export_footer_bidder_placeholder", output_language)
        # "ABN [XX XXX XXX XXX]" stays as-is in both languages -- ABN is a
        # fixed Australian legal-identifier label, not translatable content.
        registered_address = export_i18n.export_t("export_footer_registered_address", output_language)
        text = f"{bidder} | ABN [XX XXX XXX XXX] | {registered_address}"
        complete = False
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = DARK_GREY if complete else RED
    run.italic = not complete


def _build_letter_methodology(doc: Document, methodology_text: str, theme: dict | None,
                               sales_pitch_text: str | None = None, output_language: str = "en"):
    """Renders the AI-drafted "Methodology and Deliverables" section -- same
    continuous-prose, bold-**subheading** convention as every other AI section
    draft in this tool (see draft_generator.py), reusing _add_draft_paragraph
    for the bold-subheading-as-heading treatment. sales_pitch_text: the user's
    own "why choose us" text from Project Setup -- pure upside/sales copy, so
    it renders at the end of this section whenever present, independently of
    whether the methodology itself has been drafted yet."""
    theme = theme or _theme_colours(None)
    methodology_text = (methodology_text or "").strip()
    eyebrow = export_i18n.export_t("export_eyebrow_why_choose_us", output_language)
    if not methodology_text:
        _add_placeholder_paragraph(
            doc, export_i18n.export_t("export_no_methodology", output_language),
        )
        _add_pull_quote_box(doc, sales_pitch_text, theme, eyebrow=eyebrow)
        return
    for para_text in [p.strip() for p in methodology_text.split("\n\n") if p.strip()]:
        _add_draft_paragraph(doc, para_text, theme)
    _add_pull_quote_box(doc, sales_pitch_text, theme, eyebrow=eyebrow)


def _add_letter_body_text(doc: Document, text: str, placeholder: str, theme: dict | None = None):
    text = (text or "").strip()
    if not text:
        _add_placeholder_paragraph(doc, placeholder)
        return
    for para in text.split("\n\n"):
        if para.strip():
            _add_draft_paragraph(doc, para.strip(), theme)


def _build_letter_scope_of_work(doc: Document, scope_items: list, output_language: str = "en"):
    if not scope_items:
        _add_placeholder_paragraph(doc, export_i18n.export_t("export_no_scope_items", output_language))
        return
    for i, item in enumerate(scope_items, start=1):
        doc.add_heading(f"2.{i} {item.title}", level=2)
        if item.tasks:
            for task in item.tasks:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(str(task))
        else:
            _add_placeholder_paragraph(
                doc, export_i18n.export_t("export_no_tasks_for_item", output_language, item=item.title.upper()))


def _build_letter_team(doc: Document, resource_plan: list, personnel_photos: dict, theme: dict,
                       output_language: str = "en"):
    """Built from the SAME resourcing plan (Team & Resourcing tab) and the
    same "include in proposal" ticks as the Large Scope pack's Key Personnel
    profiles -- see resourcing.letter_team_entries -- rather than a separate,
    disconnected team list. A discipline's support members (added under a
    lead, each with their own custom title -- e.g. "Ryan Swagemakers, Bridge
    Engineer" under "Mat Williams, Structural") render indented directly
    under their lead, with a "->" marker, so the reporting relationship shown
    in the app's org chart preview is still visible here as plain text, even
    though this pack doesn't embed the chart image itself."""
    from modules.resourcing import letter_team_entries
    from modules.team_bios import _strip_value_prefix

    entries = letter_team_entries(resource_plan or [])
    if not entries:
        _add_placeholder_paragraph(
            doc, export_i18n.export_t("export_no_team_members", output_language),
        )
        return

    personnel_photos = personnel_photos or {}
    for entry in entries:
        name = entry["name"]
        indent = entry["indent"]

        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        photo_cell, text_cell = table.rows[0].cells
        photo_cell.width = Cm(2.2) if indent else Cm(2.6)
        text_cell.width = Cm(12.6) if indent else Cm(13.4)

        photo_bytes = _photo_for(personnel_photos, entry.get("assignment"), name) if name else None
        if photo_bytes:
            try:
                photo_cell.paragraphs[0].add_run().add_picture(io.BytesIO(photo_bytes), width=Cm(2.0 if indent else 2.4))
            except Exception:
                pass

        left_indent = Cm(0.9) if indent else None
        name_p = text_cell.paragraphs[0]
        if left_indent:
            name_p.paragraph_format.left_indent = left_indent
        name_run = name_p.add_run(("↳ " if indent else "") + (name or "[NOT ASSIGNED]") + f" | {entry['role_label']}")
        name_run.bold = True
        name_run.font.color.rgb = theme["heading"] if name else RED
        if not name:
            name_run.italic = True

        value_to_project = _strip_value_prefix(entry["value_to_project"], name)
        for label, value in [
            ("Qualification", entry["qualification"]),
            ("Experience", ", ".join(x for x in (entry["years_experience"], entry["rpeq_status"]) if x)),
            (f"On this project, {name or '[name]'} will", value_to_project),
        ]:
            p = text_cell.add_paragraph()
            if left_indent:
                p.paragraph_format.left_indent = left_indent
            r1 = p.add_run(f"{label}. ")
            r1.bold = True
            if value:
                p.add_run(value)
            else:
                r2 = p.add_run("[NOT PROVIDED]")
                r2.font.color.rgb = RED
                r2.italic = True

        if entry["relevant_projects"]:
            rp = text_cell.add_paragraph()
            if left_indent:
                rp.paragraph_format.left_indent = left_indent
            rp.add_run("Relevant project experience:").bold = True
            for item in entry["relevant_projects"]:
                bp = text_cell.add_paragraph(style="List Bullet")
                if left_indent:
                    bp.paragraph_format.left_indent = left_indent + Cm(0.4)
                bp.add_run(str(item))

        doc.add_paragraph()


def _build_letter_fee_buildup(doc: Document, discipline_fee_lines: list, theme: dict,
                              output_language: str = "en"):
    """The discipline fee build-up (hours x rate, Fees & Program tab) as a
    per-discipline $ total -- same figures, same table shape as the Large
    Scope pack's "Fee summary" (_build_commercial_section), deliberately
    showing only the resulting $ total per discipline, never the underlying
    hours/rate a client shouldn't see in a client-facing document."""
    doc.add_heading(export_i18n.export_t("heading_discipline_fee_buildup", output_language), level=2)
    lines = discipline_fee_lines or []
    if not lines:
        _add_placeholder_paragraph(
            doc, "[NO FEE BUILD-UP ENTERED -- price the discipline fee table in the Fees & Program tab]",
        )
        return
    headers = ["Discipline", "Fee (excl. GST)"]
    total = 0.0
    rows = []
    for line in lines:
        amount = getattr(line, "fee_amount", 0.0) or 0.0
        total += amount
        rows.append([line.discipline, f"${amount:,.0f}" if amount else "[ENTER FEE]"])
    rows.append(["Total", f"${total:,.0f}"])
    table = _add_table(doc, headers, rows, theme=theme)
    total_row = table.rows[-1]
    for cell in total_row.cells:
        _shade_cell(cell, str(theme["accent"]))
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def format_fee_percentages(values: list[float]) -> list[str]:
    """Format a set of fee percentages so the printed figures add up to the
    printed total.

    Rounding each share to a whole number independently is what a reader
    notices: three disciplines at 33.4/33.3/33.3 print as 33/33/33 and total
    99%, and a fee table that doesn't add to 100% reads as an arithmetic
    error in a priced offer. Largest-remainder rounding assigns the leftover
    whole points to the shares with the biggest fractional parts, so the
    printed numbers are each within one point of the true value AND sum to
    the true total.

    Deliberately whole numbers rather than a decimal place: a fee split is a
    judgement, and "33.3%" implies a precision the estimate doesn't have.
    Only applied when the underlying values genuinely total ~100; a split
    that really is 95% or 110% must keep showing that, not be quietly
    normalised into looking correct."""
    values = [float(v or 0) for v in values]
    total = sum(values)
    if not values or abs(total - 100.0) > 0.5:
        # Not a 100% split (or empty) -- show each value as it is, with a
        # decimal place where rounding would otherwise hide a difference.
        return [f"{v:.1f}%".replace(".0%", "%") for v in values]

    floors = [int(v) for v in values]
    remainder = round(total) - sum(floors)
    order = sorted(range(len(values)), key=lambda i: values[i] - floors[i], reverse=True)
    for i in order[:max(0, remainder)]:
        floors[i] += 1
    return [f"{v}%" for v in floors]


def _build_letter_fee_split(doc: Document, fee_estimates: list, theme: dict,
                            output_language: str = "en"):
    """Discipline fee % breakdown -- the table that replaced the old per-scope-item
    fee table as the one that actually goes into the pack. Deliberately
    percentage-only, no $ column: the underlying $ figures live in the discipline
    fee build-up table just above (_build_letter_fee_buildup), and the Fees &
    Program tab's "Discipline fee split" section is where the user edits this
    percentage breakdown (seeded either from that $ build-up or from the
    benchmark/AI buttons) before it's exported."""
    from modules.fee_estimation_engine import INDICATIVE_NOTE
    doc.add_heading(export_i18n.export_t("heading_fee_split_by_discipline", output_language), level=2)
    warn = doc.add_paragraph()
    run = warn.add_run(INDICATIVE_NOTE)
    run.font.bold = True
    run.font.color.rgb = RED

    headers = ["Discipline", "Fee %", "Confidence", "Source"]
    printed = format_fee_percentages([e.fee_percentage for e in fee_estimates])
    rows = [[e.discipline, pct, e.confidence, e.source]
            for e, pct in zip(fee_estimates, printed)]
    total_pct = sum(e.fee_percentage or 0 for e in fee_estimates)
    rows.append(["Total", format_fee_percentages([total_pct])[0], "-", "-"])

    table = _add_table(doc, headers, rows, theme=theme)
    total_row = table.rows[-1]
    for cell in total_row.cells:
        _shade_cell(cell, str(theme["accent"]))
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _letter_program_week_grid(doc: Document, program_schedule: dict, week_labels: list,
                              theme: dict):
    """The original week-by-scope-item shaded grid. Still here as the fallback
    for when image generation fails: a program the reader can follow beats a
    missing section, and the grid is the one form of it that needs nothing
    but python-docx."""
    headers = ["Scope Item"] + list(week_labels)
    rows = [[title] + ["" for _ in week_labels] for title in program_schedule]
    table = _add_table(doc, headers, rows, theme=theme)
    for row_index, (title, active_weeks) in enumerate(program_schedule.items(), start=1):
        for col_index, active in enumerate(active_weeks, start=1):
            if active:
                _shade_cell(table.rows[row_index].cells[col_index], str(theme["accent"]))


def _letter_program_native_table(doc: Document, model, theme: dict):
    """The formal-table style as a REAL Word table rather than a picture.

    The whole appeal of the formal-table option is that it is conservative
    and plain -- which also means it is the one style a client's document
    controller might want to re-type a row of. A picture would take that
    away for nothing."""
    def _week_text(week: int) -> str:
        label = (model.week_labels[week - 1] if week - 1 < len(model.week_labels)
                 else f"Wk {week}")
        date_text = model.week_dates[week - 1] if week - 1 < len(model.week_dates) else ""
        return f"{label} - {date_text}" if date_text else label

    rows = [
        [item.label or "[UNTITLED SCOPE ITEM]", _week_text(item.start_week),
         _week_text(item.end_week), f"{item.weeks} week{'s' if item.weeks != 1 else ''}"]
        for item in model.items
    ]
    _add_table(doc, ["Scope Item", "Commence", "Complete", "Duration"], rows, theme=theme)
    if model.start_date_text:
        note = doc.add_paragraph()
        run = note.add_run(
            f"Program anchored to an anticipated commencement of {model.start_date_text} "
            f"-- dates shift with the actual award date."
        )
        run.italic = True
        run.font.size = Pt(8.5)


def _build_letter_program(doc: Document, program_schedule: dict, week_labels: list, theme: dict,
                          style: str | None = None, methodology_stages: list | None = None,
                          start_date=None, analysis=None, project_name: str = "",
                          client_name: str = "", output_language: str = "en"):
    """The Program section, in the presentation style the user chose.

    Three of the four styles render as a full-width themed image of exactly
    what the preview showed. The formal table renders as a native Word
    table, for the reason in _letter_program_native_table(). Anything that
    goes wrong falls back to the original shaded week grid rather than
    leaving the section empty."""
    if not program_schedule or not week_labels:
        _add_placeholder_paragraph(doc, export_i18n.export_t("export_no_program_entered", output_language))
        return

    try:
        from modules import program_render

        model = program_render.build_model(
            program_schedule, week_labels, methodology_stages or [], start_date, analysis,
            project_name, client_name,
        )
        resolved = program_render.effective_style(
            model, style or program_render.DEFAULT_STYLE)
        if resolved == "table":
            _letter_program_native_table(doc, model, theme)
            return
        png = program_render.render_png(model, resolved, f"#{theme['accent']}")
        if png:
            _add_full_width_image(doc, png)
            return
    except Exception as exc:  # noqa: BLE001 -- never lose the section over a picture
        print(f"[export_docx] program image: {exc}", file=sys.stderr)

    _letter_program_week_grid(doc, program_schedule, week_labels, theme)


def _build_risk_table(doc: Document, register, theme: dict | None):
    """Risk / impact / mitigation, with TBC cells kept red.

    A "TBC" mitigation is the correct output for a brief that doesn't
    discuss mitigation, and it has to LOOK unfinished -- a black "TBC" in a
    submitted table reads as a considered answer."""
    theme = theme or _theme_colours(None)
    entries = getattr(register, "entries", None) or []
    rows = [[e.risk, e.impact, e.mitigation] for e in entries]
    table = _add_table(doc, ["Risk", "Impact", "Mitigation"], rows, theme=theme)
    for row_index, entry in enumerate(entries, start=1):
        for col_index, value in enumerate((entry.risk, entry.impact, entry.mitigation)):
            if str(value).strip().upper() != "TBC":
                continue
            cell = table.rows[row_index].cells[col_index]
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RED
                run.italic = True
    note = doc.add_paragraph()
    run = note.add_run(
        "[FIRST-PASS REGISTER -- every mitigation is a commitment this firm will be held to. "
        "Confirm each one, and replace every red TBC, before submission.]"
    )
    run.font.color.rgb = RED
    run.italic = True


def _build_letter_signoff(doc: Document, sender: dict, output_language: str = "en"):
    doc.add_paragraph()
    doc.add_paragraph(export_i18n.export_t("export_letter_signoff_regards", output_language))
    doc.add_paragraph()
    doc.add_paragraph()
    name_p = doc.add_paragraph()
    name_p.add_run(
        sender.get("name") or export_i18n.export_t("export_letter_sender_placeholder", output_language)
    ).bold = True
    if sender.get("title"):
        doc.add_paragraph(sender["title"])
    if sender.get("phone"):
        doc.add_paragraph(sender["phone"])
    if sender.get("email"):
        doc.add_paragraph(sender["email"])


def _build_letter_review_checklist(doc: Document, output_language: str = "en"):
    doc.add_heading(export_i18n.export_t("heading_letter_review_checklist", output_language), level=1)
    items = [
        export_i18n.export_t(key, output_language) for key in (
            "export_letter_checklist_placeholders",
            "export_letter_checklist_fees",
            "export_letter_checklist_team",
            "export_letter_checklist_program",
            "export_letter_checklist_terms",
            "export_letter_checklist_footer",
            "export_letter_checklist_proofread",
        )
    ]
    _add_bullets(doc, items, color=RED)


# ---------------------------------------------------------------------------
# Sections
# ---------------------------------------------------------------------------

def _build_cover_page(
    doc: Document, project_info: dict, cover_image_bytes: bytes | None,
    cover_theme_image_bytes: bytes | None = None, theme: dict | None = None,
    firm: dict | None = None, output_language: str = "en",
):
    theme = theme or _theme_colours(project_info.get("proposal_theme"))

    # NOTE: this English disclaimer_text is what gets BAKED into the
    # full-bleed cover PNG below (divider_designer.render_full_bleed_cover
    # renders it as a raster image) -- localising that is a font-rendering
    # change in divider_designer.py, out of scope here. The plain-text
    # FALLBACK cover (used only if the image render fails) gets its own
    # language-aware disclaimer further down instead.
    disclaimer_text = (
        "FIRST-PASS PREPARATION PACK -- NOT SUBMISSION READY. Generated "
        f"{datetime.now().strftime('%d %B %Y')}. Every red guidance box and bracketed "
        "placeholder must be reviewed, verified, and removed before this document is used "
        "for anything beyond internal drafting."
    )

    # True full-bleed cover: one composed A4 image (theme-coloured band on top --
    # logo placeholder, date, title, subtitle, rule, "Response to Tender"/client
    # line -- with the real project photo, or the theme colour if there's no
    # photo, filling the page all the way to the bottom edge). No white space
    # anywhere on the page, matching the reference cover the user supplied.
    # Priority for the photo: a real uploaded project photo > a themed banner
    # generated from the user's chosen layout/colours (still real, generated
    # content -- not a placeholder, just no real photo behind it).
    image_bytes = cover_image_bytes or cover_theme_image_bytes
    cover_png = divider_designer.render_full_bleed_cover(
        tender_name=project_info.get("tender_name") or "Tender Response Pack",
        project_name=project_info.get("project_name", ""),
        client_name=project_info.get("client_name", ""),
        submission_date=project_info.get("submission_date", ""),
        photo_bytes=image_bytes,
        theme_name=project_info.get("proposal_theme"),
        disclaimer_text=disclaimer_text,
        # The firm's real logo, where the profile holds one -- otherwise the
        # cover keeps its [COMPANY LOGO] box, as it always has.
        logo_bytes=(firm or {}).get("logo_bytes"),
    )
    if cover_png:
        _add_full_bleed_cover_image(doc, cover_png)
        return

    # Fallback -- plain text cover, if image rendering fails for any reason.
    # No info table here either: the Client / Bidder / Project type /
    # Submission date fields were explicitly removed from the cover.
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        project_info.get("tender_name")
        or export_i18n.export_t("export_cover_fallback_title", output_language)
    )
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = theme["heading"]

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(project_info.get("project_name", ""))
    run.font.size = Pt(15)
    run.font.color.rgb = DARK_GREY

    doc.add_paragraph()
    if image_bytes:
        try:
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_para.add_run().add_picture(io.BytesIO(image_bytes), width=Cm(16))
        except Exception:
            _add_placeholder_paragraph(doc, export_i18n.export_t("export_cover_image_placeholder", output_language))
    else:
        _add_placeholder_paragraph(
            doc, export_i18n.export_t("export_cover_image_placeholder_detail", output_language))

    doc.add_paragraph()
    warn = doc.add_paragraph()
    warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = warn.add_run(export_i18n.export_t(
        "export_cover_disclaimer", output_language, date=datetime.now().strftime('%d %B %Y')))
    run.font.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RED
    doc.add_page_break()


def _add_ocr_notice(doc: Document, ocr_note: str | None):
    """Red review-note paragraph flagging OCR-derived source text -- added
    right after the cover/title page of each exported document when any of
    the brief's text came from OCR of scanned pages (see
    document_processor.OCR_VERIFY_TAG). Same red as the guidance notes, so
    it reads as 'review this before submission', not part of the content."""
    if not ocr_note:
        return
    p = doc.add_paragraph()
    run = p.add_run("OCR NOTICE -- " + ocr_note)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RED


def _add_toc(doc: Document, output_language: str = "en"):
    doc.add_heading(export_i18n.export_t("heading_toc", output_language), level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Right-click here and choose 'Update Field' (or Update Table) to generate the "
        "table of contents once the document is finalised."
    ).font.italic = True

    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    for el in (fldChar1, instrText, fldChar2, fldChar3):
        run._r.append(el)


def _build_executive_summary(doc: Document, exec_summary, project_info: dict | None, theme: dict | None,
                              differentiator_text: str | None = None, output_language: str = "en"):
    """Renders the unweighted Executive Summary as the first real content page
    after the cover/TOC -- a short warm intro, then catchy-titled, sales-forward
    blocks in a two-column magazine layout. exec_summary is an
    executive_summary.ExecutiveSummary (or None if not drafted yet).
    differentiator_text: the user's own "what sets us apart" text from Project
    Setup -- pure upside/sales copy, not something that needs a red placeholder
    wall when missing, so it renders (via _add_pull_quote_box) whenever present,
    independently of whether an AI executive summary has been drafted yet."""
    theme = theme or _theme_colours(None)
    doc.add_heading(export_i18n.export_t("heading_executive_summary", output_language), level=1)

    intro = getattr(exec_summary, "intro", "") if exec_summary else ""
    blocks = getattr(exec_summary, "blocks", None) if exec_summary else None
    eyebrow = export_i18n.export_t("export_pull_quote_eyebrow_differentiator", output_language)

    if not intro and not blocks:
        _add_placeholder_paragraph(
            doc, export_i18n.export_t("export_no_executive_summary", output_language),
        )
        doc.add_paragraph()
        _add_pull_quote_box(doc, differentiator_text, theme, eyebrow=eyebrow)
        return

    if intro:
        intro_p = doc.add_paragraph()
        intro_run = intro_p.add_run(intro)
        intro_run.font.size = Pt(11.5)
        doc.add_paragraph()

    if blocks:
        _start_columns(doc, 2)
        for block in blocks:
            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_before = Pt(8)
            title_p.paragraph_format.space_after = Pt(2)
            title_run = title_p.add_run(
                block.title or export_i18n.export_t("export_block_untitled", output_language))
            title_run.bold = True
            title_run.font.size = Pt(13)
            title_run.font.color.rgb = theme["heading"]

            body_p = doc.add_paragraph()
            body_p.paragraph_format.space_after = Pt(10)
            body_run = body_p.add_run(
                block.body or export_i18n.export_t("export_no_content_drafted", output_language))
            body_run.font.size = Pt(11)
            if not block.body:
                body_run.font.color.rgb = RED
                body_run.italic = True
        _end_columns(doc)

    _add_pull_quote_box(doc, differentiator_text, theme, eyebrow=eyebrow)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(6)
    note_run = note.add_run(export_i18n.export_t("export_unweighted_note", output_language))
    note_run.italic = True
    note_run.font.size = Pt(8)
    note_run.font.color.rgb = RED
    doc.add_paragraph()


def _build_tender_summary(doc: Document, analysis, weighting_chart_png: bytes | None, theme: dict,
                          output_language: str = "en"):
    doc.add_heading(export_i18n.export_t("heading_tender_summary", output_language), level=1)
    _add_labelled_paragraph(doc, "Project scope", analysis.project_scope or "(not extracted)")
    if analysis.client_objectives:
        doc.add_heading(export_i18n.export_t("heading_client_objectives", output_language), level=2)
        _add_bullets(doc, analysis.client_objectives)
    _add_labelled_paragraph(doc, "Submission date", analysis.submission_date or "(not stated)")
    if analysis.mandatory_requirements:
        doc.add_heading(export_i18n.export_t("heading_mandatory_requirements", output_language), level=2)
        _add_bullets(doc, analysis.mandatory_requirements)
    if analysis.deliverables:
        doc.add_heading(export_i18n.export_t("heading_deliverables", output_language), level=2)
        _add_bullets(doc, analysis.deliverables)
    if analysis.required_forms:
        doc.add_heading(export_i18n.export_t("heading_required_forms", output_language), level=2)
        _add_bullets(doc, analysis.required_forms)
    if analysis.risks:
        doc.add_heading(export_i18n.export_t("heading_risks_noted_in_brief", output_language), level=2)
        _add_bullets(doc, analysis.risks)
    if analysis.assumptions:
        doc.add_heading(export_i18n.export_t("heading_assumptions", output_language), level=2)
        _add_bullets(doc, analysis.assumptions)
    if analysis.disciplines_involved:
        _add_labelled_paragraph(doc, "Disciplines involved", ", ".join(analysis.disciplines_involved))
    if analysis.fee_cap:
        _add_labelled_paragraph(doc, "Stated fee cap / budget ceiling", analysis.fee_cap)

    if weighting_chart_png:
        doc.add_heading(export_i18n.export_t("heading_evaluation_weighting_dashboard", output_language), level=2)
        try:
            doc.add_picture(io.BytesIO(weighting_chart_png), width=Cm(15))
        except Exception:
            _add_placeholder_paragraph(doc, "[EVALUATION WEIGHTING DASHBOARD PLACEHOLDER]")

    if analysis.analysis_warnings:
        doc.add_heading(export_i18n.export_t("heading_extraction_warnings", output_language), level=2)
        _add_bullets(doc, analysis.analysis_warnings, color=RED)


def _build_compliance_matrix(doc: Document, compliance_items: list, theme: dict, output_language: str = "en"):
    doc.add_heading(export_i18n.export_t("heading_compliance_matrix", output_language), level=1)
    doc.add_paragraph(
        "Every requirement identified in the brief, mapped to a proposal section and a status. "
        "'Missing' items need user input before this pack is usable."
    )
    headers = ["ID", "Description", "Type", "Mapped Section", "Priority", "Status", "Action Required"]
    rows = [
        [i.requirement_id, i.description, i.requirement_type, i.mapped_section or "-",
         i.priority, i.status, i.user_action_required or "-"]
        for i in compliance_items
    ]
    _add_status_table(doc, headers, rows, status_col_index=5, theme=theme)


def _build_gap_analysis(doc: Document, gap_items: list, theme: dict, output_language: str = "en"):
    doc.add_heading(export_i18n.export_t("heading_gap_analysis", output_language), level=1)
    doc.add_paragraph("Risks and gaps this pack could identify automatically -- nothing here is invented.")
    headers = ["Risk", "Issue", "Impact", "Recommended Action", "Mapped Section"]
    rows = [[g.risk_level, g.issue, g.impact, g.recommended_action, g.mapped_section or "-"] for g in gap_items]
    _add_status_table(doc, headers, rows, status_col_index=0, theme=theme)


_PAGE_LIMIT_REASONS = {
    "tender_section_limit": "Stated explicitly in the brief for this section.",
    "weighted_total_limit": "Weighted share of the brief's stated total page limit.",
    "manual_override": "Manually set by the user.",
    "default_template": "No section-specific or total page limit stated anywhere in the brief; "
                         "applied the default page allocation template -- confirm this is reasonable.",
}


def _build_page_allocation_plan(doc: Document, sections: list, theme: dict, output_language: str = "en"):
    doc.add_heading(export_i18n.export_t("heading_page_allocation_plan", output_language), level=1)
    headers = ["Section", "Weighting", "Page Limit Source", "Allocated Pages", "Reason"]
    rows = [
        [s.title, f"{s.weighting:.0f}%", _friendly_source(s.page_limit_source),
         str(s.allocated_pages),
         _PAGE_LIMIT_REASONS.get(s.page_limit_source, "No section-specific or total page limit "
                                  "stated anywhere in the brief; applied the default page allocation template.")]
        for s in sections
    ]
    _add_table(doc, headers, rows, theme=theme)


# The three fee presentations the app can build, in the fixed order they
# render when more than one is ticked. Stable order matters: a proposal that
# put the same tables in a different sequence on each regeneration would be
# a nuisance to review against a previous version.
FEE_PRESENTATIONS = ("pct_split", "discipline_buildup", "scope_buildup")

FEE_NOTHING_SELECTED = "[SELECT WHICH FEE PRESENTATION TO INCLUDE -- Fee Estimate tab]"


def fee_sections(included: dict | None) -> dict:
    """Normalise the fee-presentation ticks, defaulting to what both packs
    exported before the choice existed."""
    included = included or {}
    return {
        "pct_split": bool(included.get("pct_split", True)),
        "discipline_buildup": bool(included.get("discipline_buildup", True)),
        "scope_buildup": bool(included.get("scope_buildup", False)),
    }


def _build_scope_item_fees(doc: Document, scope_item_fees: list | None, theme: dict | None,
                           heading_level: int = 2, output_language: str = "en"):
    """The scope-item fee build-up.

    Never exported by either pack until the user could tick it: its own
    on-screen caption calls it internal tracking. Ticked, it renders like the
    other two -- real figures where priced, an explicit red [ENTER FEE] where
    not, never a $0 that reads as free."""
    theme = theme or _theme_colours(None)
    doc.add_heading(export_i18n.export_t("heading_fee_by_scope_item", output_language), level=heading_level)
    rows_in = scope_item_fees or []
    if not rows_in:
        _add_placeholder_paragraph(
            doc, "[NO SCOPE ITEM FEES ENTERED -- price the scope item table in the Fee Estimate tab]",
        )
        return
    total = 0.0
    rows = []
    for item in rows_in:
        amount = getattr(item, "fee_amount", 0.0) or 0.0
        total += amount
        rows.append([
            getattr(item, "item_title", "") or "[UNTITLED SCOPE ITEM]",
            f"${amount:,.0f}" if amount else "[ENTER FEE]",
            getattr(item, "notes", "") or "",
        ])
    rows.append(["Total", f"${total:,.0f}", ""])
    table = _add_table(doc, ["Scope item", "Fee (excl. GST)", "Notes"], rows, theme=theme)
    for cell in table.rows[-1].cells:
        _shade_cell(cell, str(theme["accent"]))
        if cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _build_fee_estimate(
    doc: Document, fee_estimates: list, fee_cap_text: str | None, theme: dict,
    indicative_amounts: dict | None = None, output_language: str = "en",
):
    """indicative_amounts, if given, overrides each estimate's stored fee_amount --
    same override the Fee Estimate tab's "Indicative benchmark split" section
    already applies for its own on-screen table/Excel/chart (see
    fee_estimation_engine.fee_estimates_to_excel's matching parameter), for when
    the user typed a manual total project fee there rather than relying on the
    brief's stated fee cap. Without this, that manual total never reached the
    export, so this table showed "-" for every discipline even though the app's
    own Fee Estimate tab was showing real dollar figures."""
    from modules.fee_estimation_engine import INDICATIVE_NOTE
    doc.add_heading(export_i18n.export_t("heading_fee_estimate_by_discipline", output_language), level=1)
    warn = doc.add_paragraph()
    run = warn.add_run(INDICATIVE_NOTE)
    run.font.bold = True
    run.font.color.rgb = RED
    if fee_cap_text:
        doc.add_paragraph(f"Anchored to the brief's stated fee cap: {fee_cap_text}")
    headers = ["Discipline", "Fee %", "Indicative $", "Confidence", "Source"]
    indicative_amounts = indicative_amounts or {}

    rows = []
    total_pct = 0.0
    total_amount = 0.0
    any_amount = False
    printed = format_fee_percentages([e.fee_percentage for e in fee_estimates])
    for e, pct in zip(fee_estimates, printed):
        amount = indicative_amounts.get(e.discipline, e.fee_amount)
        total_pct += e.fee_percentage or 0
        if amount:
            total_amount += amount
            any_amount = True
        rows.append([
            e.discipline, pct,
            f"${amount:,.0f}" if amount else "-", e.confidence, e.source,
        ])
    if fee_estimates:
        rows.append([
            "Total", format_fee_percentages([total_pct])[0],
            f"${total_amount:,.0f}" if any_amount else "-", "-", "-",
        ])

    table = _add_table(doc, headers, rows, theme=theme)
    if fee_estimates:
        total_row = table.rows[-1]
        for cell in total_row.cells:
            _shade_cell(cell, str(theme["accent"]))
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


_PERSONNEL_HINTS = (
    "personnel", "key people", "key staff", "key team", "team member", "nominated team",
    "nominated personnel", "our team", "our people", "project team", "resourc",
    "organisation", "organization", "staff profile", "staff resume", "cv of",
    "curriculum vitae", "technical skills of",
)
_METHODOLOGY_HINTS = ("methodolog", "approach", "delivering the service", "our method", "delivery approach", "how we")
_EXPERIENCE_HINTS = ("relevant experience", "experience", "track record", "past performance", "project references")
_EXECUTIVE_SUMMARY_HINTS = ("executive summary",)


def _is_personnel_section(title: str) -> bool:
    t = (title or "").lower()
    return any(h in t for h in _PERSONNEL_HINTS)


def _is_executive_summary_section(title: str) -> bool:
    t = (title or "").lower()
    return any(h in t for h in _EXECUTIVE_SUMMARY_HINTS)


def _is_methodology_section(title: str) -> bool:
    t = (title or "").lower()
    return any(h in t for h in _METHODOLOGY_HINTS)


def _is_experience_section(title: str) -> bool:
    t = (title or "").lower()
    return any(h in t for h in _EXPERIENCE_HINTS)


def _build_proposal_response(
    doc: Document, sections: list, guidance_notes: dict, drafts: dict, graphics: list,
    divider_images: dict[str, bytes] | None = None, analysis=None,
    resource_plan: list | None = None, org_chart_png: bytes | None = None, theme: dict | None = None,
    personnel_photos: dict[str, bytes] | None = None,
    reference_projects: list | None = None, reference_project_photos: dict[str, bytes] | None = None,
    discipline_fee_lines: list | None = None, project_info: dict | None = None,
    team_intro=None,
    experience_intro=None,
    sales_pitch_text: str | None = None,
    program_schedule: dict[str, list[bool]] | None = None,
    program_week_labels: list[str] | None = None,
    firm: dict | None = None,
    fee_sections_included: dict | None = None,
    scope_item_fees: list | None = None,
    output_language: str = "en",
):
    firm = firm or {}
    fee_included = fee_sections(fee_sections_included)
    divider_images = divider_images or {}
    resource_plan = resource_plan or []
    reference_projects = reference_projects or []
    graphics_by_section: dict[str, list] = {}
    for g in graphics:
        graphics_by_section.setdefault(g.suggested_placement, []).append(g)

    assigned_people = any((getattr(a, "person_name", "") or "").strip() for a in resource_plan)
    local_benefit_needed = _local_benefit_signal(analysis)

    # Every brief structures its evaluation criteria differently -- sometimes Key
    # Personnel and Relevant Experience are two separate criteria, sometimes a brief
    # folds them into one ("Technical Skills of Key Team Members" with no separate
    # experience criterion anywhere). `sections` here is already whatever THIS brief's
    # own weighting/structure produced (see proposal_structure.py/weighting_engine.py),
    # so checking it directly -- rather than hardcoding a section number/title -- is
    # what makes the placement below adapt automatically brief to brief. If no section
    # in this particular proposal is experience-flavoured, project experience has
    # nowhere else to go, so it folds into the personnel section instead of silently
    # never appearing (see _build_sc1_project_experience_compact below).
    has_dedicated_experience_section = any(_is_experience_section(s.title) for s in sections)

    for section in sections:
        # Executive Summary is a fixed section (proposal_structure.FIXED_SECTIONS)
        # so it's always present here with its own page-allocation entry, but its
        # actual content is already rendered once, in full, by
        # _build_executive_summary() as the dedicated front-matter page right
        # after the cover/TOC (see build_docx). Rendering it again here -- as an
        # ordinary numbered section with its own guidance box and a second,
        # independently-drafted "Executive Summary" AI response -- printed two
        # different Executive Summaries into the same document. It still gets
        # its 1-page budget in the Page Allocation Plan table; it just doesn't
        # get a second content section down here.
        if _is_executive_summary_section(section.title):
            continue

        # 1. Full-page divider on its own page (full bleed, edge to edge).
        divider_png = divider_images.get(section.title)
        if divider_png:
            _add_full_page_image(doc, divider_png)
        else:
            doc.add_page_break()

        # 2. On the next page: the section heading, then the red guidance box
        #    (moved here, after the divider), then the content.
        doc.add_heading(f"{section.section_number}. {section.title}", level=1)

        note = guidance_notes.get(section.title)
        if note:
            _add_red_guidance_block(doc, note)
        # Length against the section's own page budget. A draft at a third of
        # its allocation reads as thin to an evaluator scoring it, and a draft
        # well over won't fit -- neither is visible without counting words by
        # hand, so the guidance note says it.
        _add_length_note(doc, section, (drafts or {}).get(section.title))

        draft = drafts.get(section.title)

        # 3. A coloured callout box carrying the section's lead statement (drawn
        #    from real draft content, never invented) -- the black-box highlight.
        if draft and (draft.draft_text or "").strip():
            lead = _lead_sentence(draft.draft_text)
            if lead:
                _add_callout_box(doc, lead, theme)

        # 4. Special content blocks that come from real structured data, not AI prose.
        # A Key Personnel section built here always reflects the LIVE resourcing
        # plan, never a stale AI draft -- see _build_personnel_profiles's docstring
        # for why that mismatch used to happen.
        is_personnel = _is_personnel_section(section.title)
        is_experience = _is_experience_section(section.title)
        # This brief has nowhere else for project experience to live -- fold a
        # compact version into the personnel section rather than dropping it.
        fold_experience_into_personnel = (
            is_personnel and not has_dedicated_experience_section and bool(reference_projects)
        )
        if is_personnel and resource_plan:
            _build_personnel_block(doc, resource_plan, org_chart_png, theme, personnel_photos, team_intro,
                                   output_language=output_language)
        if fold_experience_into_personnel:
            _build_sc1_project_experience_compact(
                doc, reference_projects, reference_project_photos, theme, experience_intro,
                output_language=output_language,
            )
        if is_experience:
            _build_reference_experience(
                doc, reference_projects, reference_project_photos, theme, experience_intro,
                output_language=output_language,
            )
            _build_personnel_project_matrix(doc, resource_plan, reference_projects, theme, output_language)
        if _is_methodology_section(section.title) and analysis is not None:
            _build_methodology_table(doc, analysis, theme, output_language)
        if _is_relationship_section(section.title):
            _build_relationship_management(doc, project_info, theme, firm, output_language)
        if _is_commercial_section(section.title):
            _build_commercial_section(doc, discipline_fee_lines, theme,
                                      program_schedule=program_schedule,
                                      program_week_labels=program_week_labels,
                                      fee_included=fee_included,
                                      scope_item_fees=scope_item_fees,
                                      output_language=output_language)
            if local_benefit_needed:
                _build_local_benefits(doc, project_info, theme, firm, output_language=output_language)
        elif _is_local_benefit_section(section.title):
            _build_local_benefits(doc, project_info, theme, firm, output_language=output_language)

        # 5. The first-pass draft body, in two columns like a real proposal --
        # skipped where structured, deterministic content already covers the
        # section (Key Personnel with people assigned, or with project experience
        # folded in above; Relevant Experience with reference projects entered),
        # since the AI free-text draft for those duplicated/could contradict the
        # real data now shown above (this was the source of Key Personnel names
        # not matching what was entered elsewhere).
        skip_ai_body = (
            (is_personnel and (assigned_people or fold_experience_into_personnel))
            or (is_experience and reference_projects)
        )
        if skip_ai_body:
            pass
        elif draft:
            if draft.draft_heading:
                doc.add_heading(draft.draft_heading, level=2)
            _start_columns(doc, 2)
            body = [p.strip() for p in (draft.draft_text or "").split("\n\n") if p.strip()]
            if body:
                for para_text in body:
                    _add_draft_paragraph(doc, para_text, theme)
            else:
                doc.add_paragraph("[NO DRAFT BODY -- generate a draft for this section]")
            _end_columns(doc)
        else:
            # .italic on a Paragraph is a silent no-op -- python-docx only
            # honours it on a Run -- so this was the one placeholder in the
            # document rendering as ordinary black body text, indistinguishable
            # from real drafted content at a glance.
            _add_placeholder_paragraph(
                doc, "[NO DRAFT GENERATED YET -- run Draft Responses for this section]")

        # 5b. The sales pitch, pinned to the very end of the Methodology
        # section's content (after the table and the AI draft body, not
        # alongside them) -- same "pure upside copy, render whenever present"
        # treatment as the differentiator in _build_executive_summary.
        if _is_methodology_section(section.title):
            _add_pull_quote_box(doc, sales_pitch_text, theme, eyebrow="Why choose us")

        # 6. Remaining graphic placeholders (full width).
        section_graphics = graphics_by_section.get(section.title, [])
        if divider_png:
            section_graphics = [g for g in section_graphics if g.graphic_type != "Section Divider Image"]
        # The Key Personnel section above already carries its own org chart
        # placeholder (see _build_personnel_block), so a second one listed here
        # under "Graphics for this section" would just be a duplicate.
        if is_personnel:
            section_graphics = [g for g in section_graphics if "org" not in g.graphic_title.lower()]
        if section_graphics:
            # TODO A3 i18n: minor level-3 sub-heading, out of scope for this
            # pass (major top-level/level-2 structural headings only).
            doc.add_heading("Graphics for this section", level=3)
            for g in section_graphics:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{g.graphic_title} ({g.graphic_type}) -- {g.purpose} ").italic = False
                if g.status != "Generated":
                    p.add_run(g.placeholder_text or f"[{g.graphic_title.upper()} PLACEHOLDER]").font.color.rgb = RED


_BOLD_MARKDOWN_RE = re.compile(r"\*\*(.+?)\*\*")


def _add_markdown_runs(paragraph, text: str):
    """Adds runs to `paragraph` from `text`, rendering **bold** markdown spans as
    real bold runs instead of leaving the literal asterisks in the document."""
    pos = 0
    for m in _BOLD_MARKDOWN_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_draft_paragraph(doc: Document, text: str, theme: dict | None = None):
    """Renders one paragraph of AI-drafted body text, honouring the **bold**
    markdown the drafting prompt now asks for. A paragraph that is ENTIRELY one
    bold span (e.g. "**Site investigation and staging**" with nothing else) is
    treated as an in-body subheading -- bold, theme-coloured, a touch larger,
    with extra spacing -- rather than plain inline emphasis (which just uses a
    bold run, no colour/size change, mid-sentence)."""
    stripped = text.strip()
    whole_match = re.fullmatch(r"\*\*(.+)\*\*", stripped)
    p = doc.add_paragraph()
    if whole_match:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(whole_match.group(1))
        run.bold = True
        run.font.size = Pt(12.5)
        if theme:
            run.font.color.rgb = theme["heading"]
    else:
        _add_markdown_runs(p, text)
    return p


def _lead_sentence(text: str, max_len: int = 320) -> str:
    """The section's opening statement for the callout box -- the first sentence
    (or two) of the real draft, trimmed. Never invented content. The callout box
    is already rendered bold, so **markdown** bold markers from the draft (e.g. a
    leading subheading) are stripped rather than shown as literal asterisks."""
    text = " ".join((text or "").split())
    text = _BOLD_MARKDOWN_RE.sub(r"\1", text)
    if not text:
        return ""
    out = ""
    for sentence in _split_sentences(text):
        if out and len(out) + len(sentence) > max_len:
            break
        out = (out + " " + sentence).strip()
        if len(out) >= max_len:
            break
    return out[:max_len].strip()


def _split_sentences(text: str) -> list[str]:
    import re
    parts = re.split(r"(?<=[.!?])\s+", text)
    return [p for p in parts if p.strip()]


def _add_callout_box(doc: Document, text: str, theme: dict | None):
    """A single-cell coloured box with white text -- the highlighted lead-in
    statement, like the black box in the example proposal."""
    theme = theme or _theme_colours(None)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _shade_cell(cell, str(theme["primary"]))
    _set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    p = cell.paragraphs[0]
    # A short accent bar above the text.
    bar = p
    run = bar.add_run(text)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(12)
    run.font.bold = True
    doc.add_paragraph()


def _tint_colour(rgb: RGBColor, amount: float) -> RGBColor:
    """Lightens an RGBColor toward white by `amount` (0 = unchanged, 1 = white)
    -- used to derive the pull-quote box's light panel from the theme's dark
    heading colour, so it stays legible under every theme (including
    Minimalist, whose "primary" is itself a light background colour)."""
    r, g, b = rgb[0], rgb[1], rgb[2]
    return RGBColor(
        int(r + (255 - r) * amount),
        int(g + (255 - g) * amount),
        int(b + (255 - b) * amount),
    )


def _add_pull_quote_box(doc: Document, text: str, theme: dict | None, eyebrow: str | None = None):
    """The differentiator (end of Executive Summary) / sales pitch (end of
    Methodology) pull-quote card -- deliberately different from
    _add_callout_box's single-colour band, which is already used everywhere
    for every section's auto-extracted lead sentence, so this reads as
    distinct, stand-out sales copy rather than blending in as "just another
    callout". A solid tab on the left carries a large quote mark; a light
    tinted panel on the right carries an optional small uppercase eyebrow
    label and the bold body text. Built from theme["heading"] (not
    theme["primary"]) so the tab stays a legible dark fill under every theme,
    including Minimalist, whose "primary" is a light background colour, not a
    text-safe dark one. text is never invented -- this is only ever the
    user's own typed differentiator/sales-pitch text, so it's skipped
    entirely when empty rather than showing a placeholder."""
    text = (text or "").strip()
    if not text:
        return
    theme = theme or _theme_colours(None)
    tab_colour = theme["heading"]
    panel_colour = _tint_colour(tab_colour, 0.92)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    left_w, right_w = Cm(2.4), Cm(14.2)
    for row in table.rows:
        for cell, w in zip(row.cells, (left_w, right_w)):
            cell.width = w
    for col, w in zip(table.columns, (left_w, right_w)):
        col.width = w

    left_cell, right_cell = table.rows[0].cells
    _shade_cell(left_cell, str(tab_colour))
    _set_cell_margins(left_cell, top=200, bottom=200, left=100, right=100)
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    mark_p = left_cell.paragraphs[0]
    mark_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mark_run = mark_p.add_run("“")
    mark_run.font.size = Pt(44)
    mark_run.font.bold = True
    mark_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    _shade_cell(right_cell, str(panel_colour))
    _set_cell_margins(right_cell, top=200, bottom=200, left=260, right=260)
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    first_p = right_cell.paragraphs[0]
    if eyebrow:
        first_p.paragraph_format.space_after = Pt(5)
        eyebrow_run = first_p.add_run(eyebrow.upper())
        eyebrow_run.font.size = Pt(9.5)
        eyebrow_run.font.bold = True
        eyebrow_run.font.color.rgb = theme["accent"]
        body_p = right_cell.add_paragraph()
    else:
        body_p = first_p
    body_run = body_p.add_run(text)
    body_run.font.size = Pt(12.5)
    body_run.font.bold = True
    body_run.font.color.rgb = tab_colour
    doc.add_paragraph()


def _build_team_intro(doc: Document, team_intro, theme: dict | None):
    """Renders the sales-forward "why this team" introduction at the very start
    of Key Personnel, before the org chart / pen pics -- a catchy heading, a
    couple of short paragraphs (honouring **bold** markdown on named real past
    projects, same convention as the drafted section prose), and a bold italic
    closing pull-quote. team_intro is a team_intro.TeamIntro (or None if not
    drafted yet, in which case this renders nothing -- unlike other blocks,
    this one is pure upside/sales copy, not something that needs a red
    placeholder wall when missing)."""
    theme = theme or _theme_colours(None)
    heading = getattr(team_intro, "heading", "") if team_intro else ""
    paragraphs = getattr(team_intro, "paragraphs", None) if team_intro else None
    pullquote = getattr(team_intro, "pullquote", "") if team_intro else ""

    if not heading and not paragraphs:
        return

    if heading:
        heading_p = doc.add_paragraph()
        heading_p.paragraph_format.space_after = Pt(6)
        hrun = heading_p.add_run(heading)
        hrun.bold = True
        hrun.font.size = Pt(16)
        hrun.font.color.rgb = theme["heading"]

    for para_text in paragraphs or []:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        _add_markdown_runs(p, para_text)

    if pullquote:
        quote_p = doc.add_paragraph()
        quote_p.paragraph_format.space_before = Pt(4)
        quote_p.paragraph_format.space_after = Pt(14)
        qrun = quote_p.add_run(pullquote)
        qrun.bold = True
        qrun.italic = True
        qrun.font.size = Pt(12.5)
        qrun.font.color.rgb = theme["primary"]
    else:
        doc.add_paragraph()


def _build_personnel_block(
    doc: Document, resource_plan: list, org_chart_png: bytes | None, theme: dict | None,
    personnel_photos: dict[str, bytes] | None = None, team_intro=None,
    output_language: str = "en",
):
    """Key Personnel built from the ACTUAL resourcing plan the user entered (real
    assigned names/roles) plus the generated org chart -- not AI-invented names.

    Renders the team-introduction sales copy first (if drafted), then the org
    chart, then a numbered profile per key person (see
    _build_personnel_profiles) -- Project Director, Project Manager, Design
    Manager, then discipline leads, matching resourcing.personnel_profile_order().
    There is deliberately no separate "Nominated personnel" summary table any
    more: it duplicated the same names shown in the org chart and profiles
    below, using up a page for information already presented twice."""
    theme = theme or _theme_colours(None)
    _build_team_intro(doc, team_intro, theme)
    doc.add_heading(export_i18n.export_t("heading_org_chart", output_language), level=2)
    # The generated chart goes in as a FIRST PASS, and the red note below it
    # stays. Both matter. Previously org_chart_png was accepted here and
    # deliberately dropped, on the reasoning that the finished chart is built
    # in the companion PowerPoint (assets/org_chart_template.pptx) and pasted
    # in by hand -- which is still true, and still what the note says. But
    # dropping it meant a user who had assigned their whole team and watched
    # the chart render in the Team & Resourcing tab got a blank space with a
    # red instruction in the exported pack: work they had done, not shown.
    # An accurate first-pass image that says "replace me with the PowerPoint
    # version" is strictly better than an empty box that says the same thing.
    if org_chart_png:
        try:
            _add_full_width_image(doc, org_chart_png)
        except Exception:
            # An unreadable PNG must never take the whole export down -- the
            # note below still tells the user what to do.
            pass
    _add_placeholder_paragraph(
        doc,
        export_i18n.export_t("export_org_chart_firstpass_note", output_language) if org_chart_png else
        export_i18n.export_t("export_org_chart_insert_note", output_language),
    )

    _build_personnel_profiles(doc, resource_plan, personnel_photos, theme, output_language)


def _build_personnel_profiles(
    doc: Document, resource_plan: list, personnel_photos: dict[str, bytes] | None, theme: dict | None,
    output_language: str = "en",
):
    """One numbered profile per key person, in a fixed, deterministic order --
    Project Director, Project Manager, Design Manager, then discipline leads
    (resourcing.personnel_profile_order()) -- built entirely from the live
    resourcing plan at export time. This replaces relying on the AI-drafted
    free-text body for this section: that text was drafted in the Draft
    Responses step, which runs BEFORE the Team & Resourcing tab where names
    actually get assigned, so it could go stale or reference names pulled
    from uploaded company material instead of the real nominated team. Reading
    resource_plan fresh here means the names in this section always match
    whatever's currently entered in the Team & Resourcing tab."""
    from modules.resourcing import personnel_profiles_deduped

    theme = theme or _theme_colours(None)
    personnel_photos = personnel_photos or {}
    all_people = personnel_profiles_deduped(resource_plan or [])
    if not all_people:
        _add_placeholder_paragraph(doc, "[NO PERSONNEL ASSIGNED YET -- assign names in the Team & Resourcing tab]")
        return

    # Only ticked profiles make the cut -- see ResourceAssignment.include_in_proposal
    # and the "Include in proposal" checkbox next to each pen pic in the Team &
    # Resourcing tab. This is a page-space decision, not a staffing one: an
    # unticked person is still on the job (still in the org chart/fee build-up),
    # they just don't get a full photo + write-up profile here. Defaults to True,
    # so nobody drops out of an existing proposal until the user (or the AI
    # "suggest inclusion" recommendation) actually unticks them.
    people = [e for e in all_people if getattr(e["assignment"], "include_in_proposal", True)]
    if not people:
        _add_placeholder_paragraph(
            doc,
            "[NO KEY PERSONNEL ARE TICKED FOR INCLUSION -- tick at least the project "
            "leadership (Project Director/Manager/Design Manager) in the Team & "
            "Resourcing tab]",
        )
        return

    doc.add_heading(export_i18n.export_t("heading_key_personnel_profiles", output_language), level=2)
    for i, entry in enumerate(people, start=1):
        name = (entry.get("name") or "").strip()
        role_label = ", ".join(entry.get("roles") or [])

        heading = doc.add_paragraph()
        run = heading.add_run(f"{i}. {role_label}")
        run.bold = True
        run.font.size = Pt(12.5)
        run.font.color.rgb = theme["heading"]

        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        photo_cell, text_cell = table.rows[0].cells
        photo_cell.width = Cm(3.0)
        text_cell.width = Cm(13.0)

        photo_bytes = _photo_for(personnel_photos, entry.get("assignment"), name) if name else None
        if photo_bytes:
            try:
                photo_cell.paragraphs[0].add_run().add_picture(io.BytesIO(photo_bytes), width=Cm(2.8))
            except Exception:
                photo_bytes = None
        if not photo_bytes:
            _add_photo_placeholder_box(photo_cell, output_language)

        name_p = text_cell.paragraphs[0]
        name_run = name_p.add_run(name or "[INSERT KEY PERSONNEL NAME]")
        name_run.bold = True
        name_run.font.size = Pt(11.5)
        if not name:
            name_run.font.color.rgb = RED
            name_run.italic = True

        _add_personnel_field_line(text_cell, "Qualification", entry.get("qualification", ""),
                                   "[INSERT QUALIFICATION]")
        _add_personnel_field_line(text_cell, "RPEQ / registration status", entry.get("rpeq_status", ""),
                                   "[CONFIRM REGISTRATION STATUS AND NUMBER]")
        _add_personnel_field_line(text_cell, "Years of experience", entry.get("years_experience", ""),
                                   "[INSERT YEARS OF EXPERIENCE FOR CV ATTACHMENT]")

        value_p = text_cell.add_paragraph()
        vr = value_p.add_run(f"On this project, {name or '[name]'} will: ")
        vr.bold = True
        # Strip a duplicated "On this project, <name> will" opening at RENDER time -- the
        # last line of defence, catching it regardless of whether the value came from the
        # CV extraction, was typed by hand, or is old data (the run above already prints it).
        from modules.team_bios import _strip_value_prefix
        value_to_project = _strip_value_prefix((entry.get("value_to_project", "") or "").strip(), name)
        if value_to_project:
            value_p.add_run(value_to_project)
        else:
            r = value_p.add_run("[INSERT PROJECT-SPECIFIC DETAIL]")
            r.font.color.rgb = RED
            r.italic = True

        relevant_projects = entry.get("relevant_projects") or []
        if relevant_projects:
            rpp = text_cell.add_paragraph()
            rpp.add_run("Relevant project experience:").bold = True
            for item in relevant_projects:
                bp = text_cell.add_paragraph(style="List Bullet")
                bp.add_run(str(item))

        local_experience = entry.get("local_experience") or []
        if local_experience:
            lp = text_cell.add_paragraph()
            lp.add_run("Local district experience:").bold = True
            for item in local_experience:
                bp = text_cell.add_paragraph(style="List Bullet")
                bp.add_run(str(item))

        doc.add_paragraph()


def _add_personnel_field_line(cell, label: str, value: str, placeholder: str):
    p = cell.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    value = (value or "").strip()
    if value:
        p.add_run(value)
    else:
        r2 = p.add_run(placeholder)
        r2.font.color.rgb = RED
        r2.italic = True


def _add_photo_placeholder_box(cell, output_language: str = "en") -> None:
    """A shaded, empty photo slot -- as close to a one-click 'drop a headshot
    here' affordance as a static .docx can offer: click the cell, delete the
    placeholder text, then Insert > Pictures. Used for both key-personnel
    profiles and reference-project photos whenever no real image was
    uploaded, so a gap here is obvious rather than silently blank."""
    _shade_cell(cell, "E9E9EC")
    _set_cell_margins(cell, top=260, bottom=260, left=100, right=100)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(export_i18n.export_t("export_photo_placeholder", output_language))
    r.font.color.rgb = RED
    r.italic = True
    r.font.size = Pt(9)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(export_i18n.export_t("export_photo_placeholder_instruction", output_language))
    r2.font.size = Pt(7)
    r2.font.color.rgb = DARK_GREY
    r2.italic = True


METHODOLOGY_PLACEHOLDER_NOTE = (
    "[INSERT METHODOLOGY TABLE -- generate it in the app and paste the finished "
    "PowerPoint table here]"
)


def _build_methodology_table(doc: Document, analysis, theme: dict | None, output_language: str = "en"):
    """Deliberately NOT auto-building a one-column-per-scope-item methodology table
    here, and -- as of this revision -- deliberately NOT embedding a first-pass
    image of the reviewed stage grid either. That image used to sit above this
    note, but the owner wants the Word pack's body to stay clean: the finished
    methodology table belongs in PowerPoint (see modules/methodology_pptx.py,
    exported alongside this document from the Export tab -- four presentation
    styles, all built from the same reviewed stage grid the preview shows) and
    is pasted in by hand once finished. This function now does exactly one
    thing: leave the single red placeholder note marking where it goes."""
    theme = theme or _theme_colours(None)
    doc.add_heading(export_i18n.export_t("heading_methodology_summary", output_language), level=2)
    _add_placeholder_paragraph(doc, METHODOLOGY_PLACEHOLDER_NOTE)


def _build_experience_intro(doc: Document, experience_intro, theme: dict | None):
    """Renders the short sales-forward intro paragraph at the start of project
    experience -- naming the strongest 2-4 comparable reference projects and why
    they prove this firm can deliver the brief, in place of a generic "selected
    past projects" note. experience_intro is an experience_intro.ExperienceIntro
    (or None if not drafted yet, in which case the caller falls back to its own
    default note -- see _build_sc1_project_experience_compact and
    _build_reference_experience)."""
    theme = theme or _theme_colours(None)
    paragraph_text = (getattr(experience_intro, "paragraph", "") if experience_intro else "").strip()
    if not paragraph_text:
        return False
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    _add_markdown_runs(p, paragraph_text)
    return True


def _build_sc1_project_experience_compact(
    doc: Document, reference_projects: list, reference_project_photos: dict[str, bytes] | None, theme: dict | None,
    experience_intro=None, output_language: str = "en",
):
    """A compact, space-conscious companion to _build_reference_experience below --
    used ONLY when this brief's own section structure has folded Key Personnel and
    Relevant Experience into a single evaluation criterion (see
    has_dedicated_experience_section in _build_proposal_response), so project
    experience would otherwise have nowhere in the export to appear at all.

    Same visual language as the pen-pic profiles rendered immediately above this
    (_build_personnel_profiles): small photo left, text right, one project per row,
    reusing the exact real fields reference_projects.py already drafts (description,
    relevance_text) rather than SC2's wider photo-on-top cards, which cost more
    vertical space than a page-limited combined section can spare. Nothing here is
    AI-invented -- description/relevance_text are drafted once in Upload Documents
    (tab 2) and reused here and in SC2 alike; a project missing either field shows
    the same red placeholder convention as everywhere else in this tool."""
    theme = theme or _theme_colours(None)
    if not reference_projects:
        return

    reference_project_photos = reference_project_photos or {}
    doc.add_heading(export_i18n.export_t("heading_relevant_experience_compact", output_language), level=2)
    if not _build_experience_intro(doc, experience_intro, theme):
        note = doc.add_paragraph()
        note.add_run(
            "Selected past projects most relevant to this brief's scope, drawn from the "
            "firm's project reference library."
        ).italic = True

    for project in reference_projects:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        photo_cell, text_cell = table.rows[0].cells
        photo_cell.width = Cm(5.2)
        text_cell.width = Cm(11.3)

        photo_bytes = _photo_for(reference_project_photos, project, project.title)
        if photo_bytes:
            try:
                photo_cell.paragraphs[0].add_run().add_picture(io.BytesIO(photo_bytes), width=Cm(5.0))
            except Exception:
                photo_bytes = None
        if not photo_bytes:
            _add_photo_placeholder_box(photo_cell, output_language)

        title_p = text_cell.paragraphs[0]
        title_text = project.title + (f" | {project.client}" if project.client else "")
        tr = title_p.add_run(title_text)
        tr.bold = True
        tr.font.size = Pt(12.5)
        tr.font.color.rgb = theme["heading"]

        desc_p = text_cell.add_paragraph()
        desc_p.paragraph_format.space_before = Pt(4)
        description = (project.description or "").strip()
        if description:
            desc_p.add_run(description).font.size = Pt(11)
        else:
            r = desc_p.add_run(
                "[NO DESCRIPTION DRAFTED -- draft/review this reference project in Upload Docs]"
            )
            r.font.color.rgb = RED
            r.italic = True
            r.font.size = Pt(11)

        rel_p = text_cell.add_paragraph()
        rel_p.paragraph_format.space_before = Pt(6)
        rel_label = rel_p.add_run("Relevance to project: ")
        rel_label.bold = True
        rel_label.font.size = Pt(11)
        relevance = (project.relevance_text or "").strip()
        if relevance:
            rel_p.add_run(relevance).font.size = Pt(11)
        else:
            r = rel_p.add_run(
                "[NO RELEVANCE DRAFTED -- draft/review this reference project in Upload Docs]"
            )
            r.font.color.rgb = RED
            r.italic = True
            r.font.size = Pt(11)

        doc.add_paragraph()
        doc.add_paragraph()


# ---------------------------------------------------------------------------
# Relevant Experience -- reference-project profiles + key personnel cross-reference
# ---------------------------------------------------------------------------

def _build_reference_experience(
    doc: Document, reference_projects: list, reference_project_photos: dict[str, bytes] | None, theme: dict | None,
    experience_intro=None, output_language: str = "en",
):
    """Section 2 project experience, built from structured reference_projects
    entries (title/client/revised description/relevance/personnel, each with an
    optional photo) instead of dumping the raw uploaded project-reference text
    into the section. See modules/reference_projects.py for how these entries
    get drafted (revised for consistency and relevance, not copy-pasted)."""
    theme = theme or _theme_colours(None)
    if not reference_projects:
        _add_placeholder_paragraph(
            doc, "[NO REFERENCE PROJECTS ENTERED -- add project references in Upload Docs, "
                 "then draft/review them there before export]",
        )
        return

    reference_project_photos = reference_project_photos or {}
    doc.add_heading(export_i18n.export_t("heading_relevant_experience", output_language), level=2)
    _build_experience_intro(doc, experience_intro, theme)
    # Two cards side by side, each sized from the page's REAL text width
    # rather than a fixed 7.8 cm. Two 7.8 cm images plus the gap between the
    # cells overflowed A4's ~17 cm text column, so Word squashed both photos
    # to fit -- visibly, and differently depending on the image.
    photo_w = _reference_photo_width(doc)
    for start in range(0, len(reference_projects), 2):
        pair = reference_projects[start:start + 2]
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        for i, cell in enumerate(table.rows[0].cells):
            if i < len(pair):
                _fill_reference_project_cell(cell, pair[i], reference_project_photos, theme, photo_w,
                                             output_language)
        doc.add_paragraph()


def _photo_for(photos: dict, obj, fallback_name: str):
    """Look a photo up by the object's stable id, falling back to its
    name/title for anything saved before ids existed."""
    photos = photos or {}
    key = (getattr(obj, "photo_id", "") or "").strip()
    if key and key in photos:
        return photos[key]
    return photos.get((fallback_name or "").strip())


def _reference_photo_width(doc: Document):
    """Half the usable text width, less the gap between the two cards."""
    sec = doc.sections[-1]
    usable = sec.page_width - sec.left_margin - sec.right_margin
    return Emu(int((usable - Cm(0.5)) / 2))


def _fill_reference_project_cell(cell, project, photos: dict[str, bytes], theme: dict,
                                 photo_width=None, output_language: str = "en"):
    photo_width = photo_width or Cm(7.8)
    photo_bytes = _photo_for(photos, project, project.title)
    p0 = cell.paragraphs[0]
    if photo_bytes:
        try:
            p0.add_run().add_picture(io.BytesIO(photo_bytes), width=photo_width)
        except Exception:
            photo_bytes = None
    if not photo_bytes:
        # A small nested table confines the placeholder shading to just the photo
        # area -- shading the outer cell directly would grey out the description/
        # relevance/personnel text below it too.
        nested = cell.add_table(rows=1, cols=1)
        nested.autofit = False
        nested.rows[0].height = Cm(3.6)
        nested.rows[0].cells[0].width = photo_width
        _add_photo_placeholder_box(nested.rows[0].cells[0], output_language)

    title_p = cell.add_paragraph()
    title_text = project.title + (f" | {project.client}" if project.client else "")
    tr = title_p.add_run(title_text)
    tr.bold = True
    tr.font.color.rgb = theme["heading"]

    desc_p = cell.add_paragraph()
    description = (project.description or "").strip()
    if description:
        desc_p.add_run(description)
    else:
        r = desc_p.add_run("[NO DESCRIPTION DRAFTED -- draft/review this reference project in Upload Docs]")
        r.font.color.rgb = RED
        r.italic = True

    rel_p = cell.add_paragraph()
    rel_p.add_run("Relevance to project: ").bold = True
    relevance = (project.relevance_text or "").strip()
    if relevance:
        rel_p.add_run(relevance)
    else:
        r = rel_p.add_run("[INSERT RELEVANCE TO THIS TENDER]")
        r.font.color.rgb = RED
        r.italic = True

    pers_p = cell.add_paragraph()
    pers_p.add_run("Personnel involved: ").bold = True
    if project.personnel_involved:
        pers_p.add_run(", ".join(project.personnel_involved))
    else:
        r = pers_p.add_run("[CONFIRM WHICH KEY PERSONNEL WORKED ON THIS PROJECT]")
        r.font.color.rgb = RED
        r.italic = True


def _build_personnel_project_matrix(doc: Document, resource_plan: list, reference_projects: list, theme: dict | None,
                                    output_language: str = "en"):
    """Cross-references Section 3's nominated key personnel against Section 2's
    reference projects -- one row per reference project, one column per assigned
    key person, ticked where that project's personnel_involved names them."""
    from modules.resourcing import personnel_profiles_deduped

    theme = theme or _theme_colours(None)
    if not reference_projects:
        return
    # One column per unique assigned person (deduped -- a person on two disciplines
    # must not produce two identical columns), same set as the profile block above --
    # and, same as that block, only people actually included in the proposal
    # (unticked personnel, e.g. no CV provided, must not appear here either).
    people = [
        e for e in personnel_profiles_deduped(resource_plan or [])
        if (e.get("name") or "").strip() and getattr(e["assignment"], "include_in_proposal", True)
    ]
    if not people:
        return

    doc.add_heading(export_i18n.export_t("heading_personnel_experience_matrix", output_language), level=2)
    note = doc.add_paragraph()
    note.add_run("Cross-reference of which nominated key personnel worked on each reference project below.").italic = True

    headers = ["Reference project"] + [e["name"] for e in people]
    rows = []
    for project in reference_projects:
        involved = {n.strip().lower() for n in project.personnel_involved}
        rows.append([project.title] + ["✓" if e["name"].strip().lower() in involved else "" for e in people])
    table = _add_table(doc, headers, rows, theme=theme)
    for row in table.rows[1:]:
        for cell in row.cells[1:]:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Relationship management (standard-text section) + commercial / local benefits
# ---------------------------------------------------------------------------

_RELATIONSHIP_HINTS = ("relationship management", "collaborat", "partnering")
_COMMERCIAL_HINTS = ("commercial", "value for money", "price", "pricing", "fee proposal")
_LOCAL_BENEFIT_HINTS = ("local benefit", "local content", "local participation")
_LOCAL_BENEFIT_SIGNAL_KEYWORDS = (
    "local benefit", "local content", "local participation", "local industry",
    "local employment", "regional benefit", "local business", "local economy",
)

_STANDARD_RELATIONSHIP_TABLE = [
    ("Partnering approach to managing services",
     "Achieve a cooperative approach and foster positive relationships through an inception "
     "meeting to reconfirm objectives and communication lines, and regular meetings/updates "
     "from our leadership team."),
    ("Collaboration with the client, playing an informed leadership role",
     "We support the client's principal and leadership team by keeping them engaged and "
     "informed, valuing their input to drive project success."),
    ("Relational decision-making on a progressive basis",
     "Our approach provides continuous, ongoing opportunities for the client's active "
     "engagement as a key contributor to the decision-making process, including workshops at "
     "critical milestones."),
    ("Regular project meetings",
     "Productive, well-managed meetings support effective project management and give the "
     "client a record of issues, actions, and decision-making, and confidence in delivery status."),
    ("Early identification of scope changes",
     "We follow a 'no surprises', best-for-project approach, raising potential variations at "
     "the earliest opportunity, with the benefits and risks of omission outlined clearly to "
     "support the client's decision-making."),
    ("Doing business in a positive way",
     "Our team maintains a best-for-project, can-do attitude, ensuring we react positively to "
     "changes required at short notice."),
    ("Performance reporting",
     "Our success depends on our client's success. We welcome feedback through formal review "
     "systems to improve service quality and ensure value for money."),
]


def _is_relationship_section(title: str) -> bool:
    t = (title or "").lower()
    return any(h in t for h in _RELATIONSHIP_HINTS)


def _is_commercial_section(title: str) -> bool:
    t = (title or "").lower()
    return any(h in t for h in _COMMERCIAL_HINTS)


def _is_local_benefit_section(title: str) -> bool:
    t = (title or "").lower()
    return any(h in t for h in _LOCAL_BENEFIT_HINTS)


def _local_benefit_signal(analysis) -> bool:
    """True if the brief itself appears to ask for local-benefit / local-content
    commitments -- checked against the extracted scope, objectives, mandatory
    requirements, and evaluation criteria text, never guessed from the project type."""
    if analysis is None:
        return False
    parts = [
        getattr(analysis, "project_scope", "") or "",
        " ".join(getattr(analysis, "client_objectives", None) or []),
        " ".join(getattr(analysis, "mandatory_requirements", None) or []),
    ]
    for c in getattr(analysis, "evaluation_criteria", None) or []:
        parts.append(f"{getattr(c, 'name', '')} {getattr(c, 'description', '')}")
    haystack = " ".join(parts).lower()
    return any(k in haystack for k in _LOCAL_BENEFIT_SIGNAL_KEYWORDS)


def _build_relationship_management(doc: Document, project_info: dict | None, theme: dict | None,
                                   firm: dict | None = None, output_language: str = "en"):
    """A standard relationship-management narrative + principles table, used as a
    more pertinent starting point than generic AI-drafted prose for this kind of
    section -- the structure and process language a firm actually uses across
    bids, not project-invented claims. Still flagged red as text to tailor and
    verify (real local staff/offices, this project's actual engagement plan)
    before submission, same discipline as every other placeholder in the tool."""
    theme = theme or _theme_colours(None)
    client = (project_info or {}).get("client_name") or "[CLIENT NAME]"

    doc.add_heading(export_i18n.export_t("heading_relationship_management", output_language), level=2)
    intro = doc.add_paragraph()
    intro.add_run(
        f"We focus on the moments that matter -- looking beyond the technical solution to foster "
        f"a united, professional relationship with {client}. By maintaining live comment "
        f"registers and prioritising timely review closure, we minimise rework and ensure "
        f"stakeholder input is captured and actioned. Proactive engagement and clear "
        f"communication are central to our relationship management approach and underpin our "
        f"proven ability to deliver projects on time."
    )
    firm = firm or {}
    leadership = (firm.get("leadership_text") or "").strip()
    if leadership:
        # The firm profile names the leadership who actually oversee delivery,
        # so say who they are instead of shipping the same anonymous paragraph
        # to every client.
        lead_p = doc.add_paragraph()
        lead_p.add_run("Leadership oversight. ").bold = True
        lead_p.add_run(leadership)
    note = doc.add_paragraph()
    r = note.add_run(
        "[STANDARD TEXT -- confirm real local staff/offices, and tailor to this project's actual "
        "engagement plan, before submission]"
    )
    r.italic = True
    r.font.color.rgb = RED

    headers = ["Principles", "Our approach"]
    rows = [[p, a] for p, a in _STANDARD_RELATIONSHIP_TABLE]
    _add_table(doc, headers, rows, theme=theme)
    doc.add_paragraph()


def _build_commercial_section(doc: Document, discipline_fee_lines: list | None, theme: dict | None,
                              program_schedule: dict[str, list[bool]] | None = None,
                              program_week_labels: list[str] | None = None,
                              fee_included: dict | None = None,
                              scope_item_fees: list | None = None,
                              output_language: str = "en"):
    """A punchier, structured commercial section -- a fee table by discipline/
    stage with a highlighted total, then short Cash flow / Contractual
    arrangements sub-sections -- instead of a single generic AI-drafted
    paragraph. Real figures only where the user has actually priced a
    discipline in the Fee Estimate tab; everything else stays a placeholder."""
    theme = theme or _theme_colours(None)
    fee_included = fee_sections(fee_included)

    # Nothing ticked must never produce a silently fee-less proposal.
    if not any(fee_included.values()):
        doc.add_heading(export_i18n.export_t("heading_fee_summary", output_language), level=2)
        _add_placeholder_paragraph(doc, export_i18n.export_t("export_fee_nothing_selected", output_language))
        doc.add_paragraph()
        _build_cash_flow(doc, discipline_fee_lines, program_schedule, program_week_labels, theme, output_language)
        _build_contractual_arrangements(doc, output_language)
        return

    if fee_included["scope_buildup"]:
        _build_scope_item_fees(doc, scope_item_fees, theme, output_language=output_language)

    if not fee_included["discipline_buildup"]:
        # The other ticked presentation(s) have rendered; the cash flow below
        # still derives from the priced build-up whether or not it is shown.
        _build_cash_flow(doc, discipline_fee_lines, program_schedule, program_week_labels, theme, output_language)
        _build_contractual_arrangements(doc, output_language)
        return

    doc.add_heading(export_i18n.export_t("heading_fee_summary", output_language), level=2)
    lines = discipline_fee_lines or []
    if not lines:
        _add_placeholder_paragraph(
            doc, "[NO FEE BUILD-UP ENTERED -- price the discipline fee table in the Fee Estimate tab]",
        )
    else:
        headers = ["Discipline / stage", "Fee (excl. GST)"]
        total = 0.0
        rows = []
        for line in lines:
            amount = getattr(line, "fee_amount", 0.0) or 0.0
            total += amount
            rows.append([line.discipline, f"${amount:,.0f}" if amount else "[ENTER FEE]"])
        rows.append(["Total", f"${total:,.0f}"])
        table = _add_table(doc, headers, rows, theme=theme)
        total_row = table.rows[-1]
        for cell in total_row.cells:
            _shade_cell(cell, str(theme["accent"]))
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_heading(export_i18n.export_t("heading_cash_flow", output_language), level=3)
    _build_cash_flow(doc, discipline_fee_lines, program_schedule, program_week_labels, theme, output_language)

    _build_contractual_arrangements(doc, output_language)


def _build_contractual_arrangements(doc: Document, output_language: str = "en") -> None:
    doc.add_heading(export_i18n.export_t("heading_contractual_arrangements", output_language), level=3)
    p = doc.add_paragraph()
    r = p.add_run(
        "[CONFIRM THE PANEL / CONTRACT AND RATES THIS FEE IS BASED ON, AND ANY SUBCONSULTANT "
        "ARRANGEMENTS (E.G. MEMORANDUM OF UNDERSTANDING / SUBCONSULTANCY AGREEMENTS)]"
    )
    r.font.color.rgb = RED
    doc.add_paragraph()


# Above this many active weeks a week-by-week cash-flow table stops being
# readable in a proposal and starts being a spreadsheet, so weeks are grouped
# into blocks. Four weeks is the natural block: it reads as a month, which is
# also the interval most of these engagements actually invoice on.
_CASH_FLOW_MAX_ROWS = 20
_CASH_FLOW_BLOCK_WEEKS = 4


def cash_flow_rows(discipline_fee_lines: list | None,
                   program_schedule: dict[str, list[bool]] | None,
                   program_week_labels: list[str] | None) -> list[tuple[str, float, float]]:
    """First-pass cash-flow profile: the priced fee spread straight-line over
    the weeks the program actually has work in.

    Returns [(period_label, amount_this_period, cumulative)], or [] when
    either input is missing -- the caller then keeps the placeholder rather
    than inventing a profile. Straight-line is a deliberate choice: it is
    derived arithmetic over two things the user entered, not a guess at how
    their effort is really shaped, and the exported note says exactly that.

    Weeks with no scope item active carry no fee, so a program with a gap in
    the middle shows that gap rather than smearing money across it.
    """
    total = sum((getattr(line, "fee_amount", 0.0) or 0.0) for line in (discipline_fee_lines or []))
    if total <= 0 or not program_schedule or not program_week_labels:
        return []

    week_count = len(program_week_labels)
    active = []
    for index in range(week_count):
        if any(
            index < len(weeks) and bool(weeks[index])
            for weeks in program_schedule.values()
        ):
            active.append(index)
    if not active:
        return []

    per_week = total / len(active)
    weekly = [per_week if index in set(active) else 0.0 for index in range(week_count)]

    # Group only when a per-week table would be unreadably long.
    if week_count > _CASH_FLOW_MAX_ROWS:
        blocks = []
        for start in range(0, week_count, _CASH_FLOW_BLOCK_WEEKS):
            end = min(start + _CASH_FLOW_BLOCK_WEEKS, week_count)
            label = (f"{program_week_labels[start]} - {program_week_labels[end - 1]}"
                     if end - 1 > start else program_week_labels[start])
            blocks.append((label, sum(weekly[start:end])))
    else:
        blocks = [(program_week_labels[i], weekly[i]) for i in range(week_count)]

    rows = []
    running = 0.0
    for label, amount in blocks:
        running += amount
        rows.append((label, amount, running))
    return rows


def _build_cash_flow(doc: Document, discipline_fee_lines: list | None,
                     program_schedule: dict[str, list[bool]] | None,
                     program_week_labels: list[str] | None, theme: dict | None,
                     output_language: str = "en"):
    """Renders the derived cash-flow profile, or the original placeholder when
    the inputs it needs don't exist yet.

    This used to be an unconditional "[INSERT PROJECT CASH FLOW PROFILE, BASED
    ON THE FEE AND PROGRAM]" -- while the app was already holding both the fee
    build-up and the week-by-week program that placeholder was telling the user
    to go and combine by hand."""
    theme = theme or _theme_colours(None)
    rows = cash_flow_rows(discipline_fee_lines, program_schedule, program_week_labels)
    if not rows:
        _add_placeholder_paragraph(doc, export_i18n.export_t("export_cash_flow_insert", output_language))
        return

    note = doc.add_paragraph()
    run = note.add_run(
        "Indicative only, derived from your fee build-up and program by spreading the total "
        "evenly across the weeks with work programmed -- refine to your real invoicing "
        "profile before submission."
    )
    run.font.bold = True
    run.font.color.rgb = RED

    total = rows[-1][2]
    table = _add_table(
        doc,
        ["Period", "Fee (excl. GST)", "Cumulative"],
        [[label, f"${amount:,.0f}", f"${cumulative:,.0f}"] for label, amount, cumulative in rows],
        theme=theme,
    )
    total_row = table.add_row()
    for cell, text in zip(total_row.cells, ["Total", f"${total:,.0f}", f"${total:,.0f}"]):
        cell.text = text
        _shade_cell(cell, str(theme["accent"]))
        cell_run = cell.paragraphs[0].runs[0]
        cell_run.font.bold = True
        cell_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()


def _build_local_benefits(doc: Document, project_info: dict | None, theme: dict | None,
                          firm: dict | None = None, output_language: str = "en"):
    """A Local Benefits sub-section, added only when the brief itself signals a
    local-benefit/local-content requirement (see _local_benefit_signal) or the
    section is explicitly a named local-benefit criterion. Structure only --
    every figure/claim is a placeholder, never invented (a firm's real local
    office split or reinvestment program is specific to that firm and bid, and
    reusing another bid's real numbers here would be worse than a gap)."""
    theme = theme or _theme_colours(None)
    firm = firm or {}
    client = (project_info or {}).get("client_name") or "the client"

    doc.add_heading(export_i18n.export_t("heading_local_benefits", output_language), level=2)
    note = doc.add_paragraph()
    r = note.add_run(export_i18n.export_t("export_local_benefits_intro_note", output_language))
    r.italic = True
    r.font.color.rgb = RED

    # Two of these four are standing firm facts -- where the firm's offices
    # are, and what it puts back into the community. Both live in the firm
    # profile now, so they render as the firm's own words instead of an
    # instruction to go and write them again for this bid. The other two are
    # genuinely per-bid (which office delivers THIS job, which local strategy
    # THIS brief names) and keep their placeholders.
    for heading_key, placeholder, real_text in [
        ("export_local_benefits_heading_resources",
         export_i18n.export_t("export_local_benefits_confirm_local", output_language),
         firm.get("offices_text")),
        ("export_local_benefits_heading_economy",
         export_i18n.export_t("export_local_benefits_describe_economy", output_language,
                              client=client.upper()),
         None),
        ("export_local_benefits_heading_strategy",
         export_i18n.export_t("export_local_benefits_reference_strategy", output_language),
         None),
        ("export_local_benefits_heading_reinvestment",
         export_i18n.export_t("export_local_benefits_confirm_reinvestment", output_language),
         firm.get("community_text")),
    ]:
        doc.add_heading(export_i18n.export_t(heading_key, output_language), level=3)
        if (real_text or "").strip():
            for para in str(real_text).split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
            confirm = doc.add_paragraph()
            cr = confirm.add_run(export_i18n.export_t("export_local_benefits_from_profile", output_language))
            cr.font.color.rgb = RED
            cr.italic = True
        else:
            p = doc.add_paragraph()
            pr = p.add_run(placeholder)
            pr.font.color.rgb = RED
    doc.add_paragraph()


def _build_review_checklist(doc: Document, sections: list, output_language: str = "en"):
    doc.add_heading(export_i18n.export_t("heading_review_checklist", output_language), level=1)
    # The marker format named in export_checklist_replace_placeholders has to
    # match what the exporters actually write in THIS language, or the
    # instruction sends someone searching the document for a string that
    # occurs nowhere in it -- see export_i18n.PLACEHOLDER_PREFIXES.
    items = [
        export_i18n.export_t(key, output_language) for key in (
            "export_checklist_delete_boxes",
            "export_checklist_replace_placeholders",
            "export_checklist_page_limits",
            "export_checklist_schedules",
            "export_checklist_personnel",
            "export_checklist_fee_cap",
            "export_checklist_graphics",
            "export_checklist_toc",
            "export_checklist_compliance_check",
            "export_checklist_submission",
        )
    ]
    _add_bullets(doc, items)


# Every placeholder marker the exporters actually write. Used to sweep the
# generated document itself for anything the compliance/draft lists don't
# already know about (see _build_user_input_list). Sourced from
# export_i18n.ALL_PLACEHOLDER_PREFIXES (Audit Round 2, Part 4) so BOTH
# English and Spanish markers are always matched, regardless of the
# project's current output_language -- a pack can contain a mix of the two
# after a mid-project language switch, and this sweep must never silently
# under-report a Spanish placeholder just because it isn't spelled in
# English.
PLACEHOLDER_MARKERS = export_i18n.ALL_PLACEHOLDER_PREFIXES
# Matched separately, on a word boundary: a bare substring test would catch
# "TBC" inside an ordinary word.
_TBC_RE = re.compile(r"\bTBC\b")


def collect_placeholders(doc: Document, limit: int = 60) -> list[str]:
    """Every red placeholder actually present in a built document.

    The User Input Required list was assembled only from the compliance
    matrix and each draft's own self-reported required_user_inputs -- which
    between them miss most of what really ends up red on the page, because
    the exporters write their own placeholders (org chart, cash flow, fee
    rows, methodology, local content, footer, ...) and nothing was reading
    them back. Sweeping the finished document is the only way to list what
    is genuinely there."""
    found: list[str] = []
    seen: set[str] = set()

    def consider(text: str) -> None:
        text = " ".join((text or "").split())
        if not text or len(text) > 300:
            return
        upper = text.upper()
        if not any(marker in upper for marker in PLACEHOLDER_MARKERS) and not _TBC_RE.search(upper):
            return
        if text in seen:
            return
        seen.add(text)
        found.append(text)

    for paragraph in doc.paragraphs:
        consider(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                consider(cell.text)
    return found[:limit]


def _build_user_input_list(doc: Document, compliance_items: list, gap_items: list, drafts: dict,
                           document_placeholders: list | None = None, output_language: str = "en"):
    doc.add_heading(export_i18n.export_t("heading_user_input_required", output_language), level=1)
    doc.add_paragraph(export_i18n.export_t("export_user_input_intro", output_language))

    seen = set()
    entries = []
    for i in compliance_items:
        if i.status in ("Missing", "User Input Required") and i.user_action_required:
            key = i.user_action_required.strip()
            if key not in seen:
                seen.add(key)
                entries.append(f"[{i.mapped_section or 'General'}] {key}")
    for draft in drafts.values():
        for req in draft.required_user_inputs:
            key = f"{draft.section_title}: {req}".strip()
            if key not in seen:
                seen.add(key)
                entries.append(f"[{draft.section_title}] {req}")

    if document_placeholders:
        doc.add_heading(export_i18n.export_t("heading_placeholders_in_document", output_language), level=2)
        _add_bullets(doc, document_placeholders, color=RED)

    _add_bullets(doc, entries or [export_i18n.export_t("export_user_input_none", output_language)])


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _set_base_styles(doc: Document, theme: dict, font_name: str = DEFAULT_FONT):
    """Applies the proposal theme's colours and the chosen body font (Arial by
    default) to body text and headings, plus a visible accent rule under each
    Heading 1 -- so the document reads as one designed piece rather than default
    Word black-on-white, and matches the divider/cover banners from the same palette."""
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    _force_east_asian_font(normal, font_name)

    sizes = {"Heading 1": 20, "Heading 2": 14, "Heading 3": 12}
    for style_name, size in sizes.items():
        s = doc.styles[style_name]
        s.font.name = font_name
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = theme["primary"] if style_name == "Heading 3" else theme["heading"]
        _force_east_asian_font(s, font_name)
    _add_bottom_border(doc.styles["Heading 1"], theme["accent"])


def _force_east_asian_font(style, font_name: str) -> None:
    """python-docx only sets the Latin font by name; set the other script slots
    too so Word doesn't silently fall back to Calibri for some characters."""
    try:
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(attr), font_name)
    except Exception:
        pass


def _add_bottom_border(style, colour: RGBColor) -> None:
    """Adds a coloured rule under every paragraph using this style (used for Heading 1,
    so each major section break has a visible accent line, echoing the divider banners)."""
    pPr = style.element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), str(colour))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_page_numbers(doc: Document) -> None:
    """Simple 'Page X of Y' centered footer -- the kind of finishing touch every real
    proposal has and a bare python-docx document doesn't get by default."""
    _write_page_number_footer(doc.sections[0].footer)


def _write_page_number_footer(footer) -> None:
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""

    def _field(instr: str):
        run = p.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = instr
        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instrText)
        run._r.append(sep)
        run._r.append(end)

    p.add_run("Page ")
    _field("PAGE")
    p.add_run(" of ")
    _field("NUMPAGES")
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = DARK_GREY


_A4_W, _A4_H = Cm(21), Cm(29.7)
_NORMAL_V, _NORMAL_H = Cm(2.0), Cm(2.2)


def _normal_margins(section) -> None:
    section.page_width, section.page_height = _A4_W, _A4_H
    section.top_margin = section.bottom_margin = _NORMAL_V
    section.left_margin = section.right_margin = _NORMAL_H


def _add_full_page_image(doc: Document, png_bytes: bytes) -> None:
    """Insert an image as a true full-bleed A4 page: a zero-margin section holds
    the image edge-to-edge, then a normal-margin section resumes for the content
    that follows. Used for the full-page section dividers."""
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width, sec.page_height = _A4_W, _A4_H
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(0)
    sec.header_distance = sec.footer_distance = Cm(0)
    # Keep page numbering off the divider page.
    try:
        sec.footer.is_linked_to_previous = False
        for p in sec.footer.paragraphs:
            p.text = ""
    except Exception:
        pass
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    try:
        p.add_run().add_picture(io.BytesIO(png_bytes), width=_A4_W)
    except Exception:
        pass
    restore = doc.add_section(WD_SECTION.NEW_PAGE)
    _normal_margins(restore)
    try:
        restore.footer.is_linked_to_previous = True
    except Exception:
        pass


def _add_full_bleed_cover_image(doc: Document, png_bytes: bytes) -> None:
    """Like _add_full_page_image, but for the cover -- the very first thing in the
    document. Reusing the existing default section (instead of add_section, which
    would insert a page BREAK and leave a blank page 1 before the image) puts the
    image on page 1 itself; a normal-margin section afterward resumes for the
    Table of Contents and everything that follows."""
    sec = doc.sections[0]
    sec.page_width, sec.page_height = _A4_W, _A4_H
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(0)
    sec.header_distance = sec.footer_distance = Cm(0)
    # Keep page numbering off the cover page (this is the section _add_page_numbers
    # wrote the "Page X of Y" footer field into, so clear it here explicitly).
    try:
        for p in sec.footer.paragraphs:
            p.text = ""
    except Exception:
        pass
    p = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    try:
        p.add_run().add_picture(io.BytesIO(png_bytes), width=_A4_W)
    except Exception:
        pass
    restore = doc.add_section(WD_SECTION.NEW_PAGE)
    _normal_margins(restore)
    # Give the restored section its OWN page-number footer rather than linking
    # back to the cover section (whose footer was just blanked above) -- since
    # the cover is section 0, "linked to previous" has nothing valid to inherit.
    try:
        restore.footer.is_linked_to_previous = False
        _write_page_number_footer(restore.footer)
    except Exception:
        pass


def _add_full_width_image(doc: Document, png_bytes: bytes) -> None:
    sec = doc.sections[-1]
    usable = sec.page_width - sec.left_margin - sec.right_margin
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(io.BytesIO(png_bytes), width=usable)


def _set_columns(section, num: int, space_twips: int = 567) -> None:
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:equalWidth"), "1")


def _start_columns(doc: Document, num: int = 2) -> None:
    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    _normal_margins(sec)
    _set_columns(sec, num)


def _end_columns(doc: Document) -> None:
    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    _normal_margins(sec)
    _set_columns(sec, 1)


def _set_cell_margins(cell, top=100, bottom=100, left=120, right=120) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _add_labelled_paragraph(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(value)


def _add_bullets(doc: Document, items: list[str], color: RGBColor | None = None):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(str(item))
        if color:
            run.font.color.rgb = color


def _add_placeholder_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.color.rgb = RED
    run.font.italic = True


def _add_length_note(doc: Document, section, draft) -> None:
    """Flags a draft that is well under or over its page allocation."""
    if draft is None:
        return
    from modules.draft_generator import LENGTH_TOLERANCE, length_verdict, target_words

    verdict = length_verdict(section, draft)
    if not verdict:
        return
    words = len((draft.draft_text or "").split())
    target = target_words(section)
    pages = getattr(section, "allocated_pages", 1)
    p = doc.add_paragraph()
    run = p.add_run(
        f"[LENGTH: this draft is about {words} words against roughly {target} for its "
        f"{pages}-page allocation -- {'well under' if verdict == 'under' else 'well over'} "
        f"budget (more than {int(LENGTH_TOLERANCE * 100)}% out). "
        + ("Expand it with real detail, or re-check the allocation."
           if verdict == "under" else
           "Cut it back, or re-check the allocation.")
    )
    run.font.color.rgb = RED
    run.italic = True


def _add_red_guidance_block(doc: Document, note):
    heading = doc.add_paragraph()
    run = heading.add_run(f"[{note.marker}] -- {note.section_title}")
    run.font.bold = True
    run.font.color.rgb = RED
    run.font.size = Pt(12)

    for label, value in [
        ("Page limit", note.page_limit_text),
        ("Evaluation weighting", note.weighting_text),
        ("Formatting requirements", note.format_requirements_text),
    ]:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}: ")
        r1.font.bold = True
        r1.font.color.rgb = RED
        r2 = p.add_run(value)
        r2.font.color.rgb = RED

    for title, items in [
        ("Brief requirements", note.brief_requirements),
        ("Recommended content", note.recommended_content),
        ("Recommended graphics", note.recommended_graphics),
        ("User actions required before submission", note.user_actions_required),
    ]:
        if items:
            p = doc.add_paragraph()
            r = p.add_run(f"{title}:")
            r.font.bold = True
            r.font.color.rgb = RED
            for item in items:
                bp = doc.add_paragraph(style="List Bullet")
                br = bp.add_run(item)
                br.font.color.rgb = RED

    doc.add_paragraph()  # spacing before the actual draft content


def _add_table(doc: Document, headers: list[str], rows: list[list[str]], theme: dict | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h) if h not in (None, "") else " "  # never empty -> always has a run to style
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        if theme:
            # Theme-coloured header row instead of Word's generic blue table style, so
            # tables match the headings/banners rather than looking bolted-on.
            _shade_cell(cell, str(theme["primary"]))
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row_values in rows:
        row = table.add_row()
        for i, value in enumerate(row_values):
            row.cells[i].text = str(value)
    return table


def _add_status_table(doc: Document, headers: list[str], rows: list[list[str]], status_col_index: int, theme: dict | None = None):
    table = _add_table(doc, headers, rows, theme=theme)
    for row_index, row_values in enumerate(rows, start=1):
        status_value = str(row_values[status_col_index])
        color_key = None
        for key in ("High", "Missing"):
            if key in status_value:
                color_key = "High"
        if not color_key:
            for key in ("Medium", "Partially"):
                if key in status_value:
                    color_key = "Medium"
        if not color_key:
            for key in ("Low", "Covered"):
                if key in status_value and "Partially" not in status_value:
                    color_key = "Low"
        if color_key:
            _shade_cell(table.rows[row_index].cells[status_col_index], RISK_SHADING[color_key])


def _shade_cell(cell, hex_color: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def _friendly_source(source: str) -> str:
    return {
        "tender_section_limit": "Stated in brief (section-specific)",
        "weighted_total_limit": "Weighted share of stated total",
        "default_template": "Default template (no limit stated)",
        "manual_override": "Manually overridden",
    }.get(source, source)
